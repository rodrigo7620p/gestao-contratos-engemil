from __future__ import annotations

from datetime import date
from io import BytesIO
from pathlib import Path

from contract_utils import today_brt
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt, RGBColor

from reports import _scaled_widths, _set_table_geometry

BASE_DIR = Path(__file__).resolve().parent
DEFAULT_TEMPLATE = BASE_DIR / "templates" / "MODELO.docx"
BURGUNDY = "5A1235"
LIGHT_GRAY = "E7E7E7"


def brl(value):
    return f"R$ {float(value or 0):,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")


def short_date(value):
    if not value:
        return "—"
    try:
        parsed = date.fromisoformat(str(value)[:10])
        return parsed.strftime("%d/%m/%Y")
    except ValueError:
        return str(value)


def set_cell_fill(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color)
    tc_pr.append(shading)


def set_cell_width(cell, inches):
    cell.width = Inches(inches)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(inches * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def style_run(run, size=11, bold=False, color="000000"):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(8 if level == 1 else 4)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.5
    run = paragraph.add_run(text)
    style_run(run, size=11, bold=True, color=BURGUNDY)
    return paragraph


def add_body(doc, text, bold=False, center=False):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.5
    style_run(paragraph.add_run(text), size=11, bold=bold)
    return paragraph


def set_paragraph_top_border(paragraph, color="000000", size="8", space="6"):
    paragraph_properties = paragraph._p.get_or_add_pPr()
    borders = paragraph_properties.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        paragraph_properties.append(borders)
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), size)
    top.set(qn("w:space"), space)
    top.set(qn("w:color"), color)
    borders.append(top)


def add_signature_text(doc, text, bold=False, border=False):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    style_run(paragraph.add_run(text), size=11, bold=bold)
    if border:
        set_paragraph_top_border(paragraph)
    return paragraph


def generate_declaration(contracts, parameters, template_path=DEFAULT_TEMPLATE):
    doc = Document(str(template_path))
    for section in doc.sections:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Mm(210)
        section.page_height = Mm(297)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), "Calibri")
    normal._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1.5
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(8)
    title.paragraph_format.space_after = Pt(0)
    title.paragraph_format.line_spacing = 1.5
    style_run(
        title.add_run("DECLARAÇÃO DE CONTRATOS FIRMADOS COM A INICIATIVA PRIVADA E A ADMINISTRAÇÃO PÚBLICA"),
        size=11, bold=True, color=BURGUNDY,
    )
    add_body(
        doc,
        "Declaramos que a Engemil – Engenharia, Empreendimentos, Manutenção e Instalações Ltda., "
        "inscrita no CNPJ nº 04.768.702/0001-70, possui os seguintes contratos firmados com a "
        "iniciativa privada e a Administração Pública, vigentes na data desta declaração:",
    )
    headers = ["Item", "Contratante e contrato", "Vigência e instrumento", "Valores"]
    widths = [0.45, 2.72, 1.72, 1.97]
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = "Table Grid"
    set_repeat_table_header(table.rows[0])
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_width(cell, widths[index])
        set_cell_fill(cell, BURGUNDY)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style_run(paragraph.add_run(header), size=11, bold=True, color="FFFFFF")
    total = 0.0
    total_remaining = 0.0
    for item, contract in enumerate(contracts, start=1):
        total += float(contract.get("current_value") or 0)
        total_remaining += float(contract.get("remaining_value") or 0)
        values = [
            str(item),
            "\n".join(filter(None, [
                f"Contratante: {contract.get('client')}" if contract.get("client") else "",
                f"Contrato: {contract.get('contract_number')}" if contract.get("contract_number") else "",
            ])),
            "\n".join(filter(None, [
                f"Início original: {short_date(contract.get('start_date'))}",
                f"Fim vigente: {short_date(contract.get('end_date'))}",
                f"Instrumento: {contract.get('current_instrument') or 'Contrato'}",
            ])),
            "\n".join([
                f"Valor atual: {brl(contract.get('current_value'))}",
                f"Remanescente: {brl(contract.get('remaining_value'))}",
            ]),
        ]
        row = table.add_row()
        prevent_row_split(row)
        for index, value in enumerate(values):
            cell = row.cells[index]
            set_cell_width(cell, widths[index])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.alignment = (
                WD_ALIGN_PARAGRAPH.CENTER if index == 0 else WD_ALIGN_PARAGRAPH.LEFT
            )
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.0
            lines = str(value or "—").splitlines()
            for line_index, line in enumerate(lines):
                if line_index:
                    paragraph.add_run().add_break()
                if ": " in line:
                    label, detail = line.split(": ", 1)
                    style_run(paragraph.add_run(f"{label}: "), size=11, bold=True)
                    style_run(paragraph.add_run(detail), size=11)
                else:
                    style_run(paragraph.add_run(line), size=11)
    _set_table_geometry(table, _scaled_widths(doc, [0.45, 2.72, 1.72, 1.97]))
    doc.add_page_break()
    add_heading(
        doc,
        'FÓRMULA PARA FINS DE ATENDIMENTO AO ITEM "D.1" DA ALÍNEA "D" DO SUBITEM 11.1 '
        "DO ITEM 11 DO ANEXO VII-A DA IN SEGES/MP Nº 05/2017",
    )
    add_body(
        doc,
        "Declaramos que 1/12 (um doze avos) dos contratos firmados com a Administração Pública "
        "e/ou com a iniciativa privada vigentes na data de apresentação da proposta não é "
        "superior ao Patrimônio Líquido da empresa.",
    )
    equity = float(parameters.get("equity_value") or 0)
    index_total = equity * 12 / total if total else 0
    index_remaining = equity * 12 / total_remaining if total_remaining else 0
    add_heading(doc, "Cálculo sobre o valor total dos contratos", level=2)
    add_body(
        doc,
        f"{brl(equity)} × 12 ÷ {brl(total)} = {index_total:.2f} "
        f"({'ATENDE' if index_total >= 1 else 'NÃO ATENDE'} ao resultado mínimo de 1,00).",
        bold=True, center=True,
    )
    add_heading(doc, "Cálculo sobre o valor remanescente dos contratos", level=2)
    add_body(
        doc,
        f"{brl(equity)} × 12 ÷ {brl(total_remaining)} = {index_remaining:.2f} "
        f"({'ATENDE' if index_remaining >= 1 else 'NÃO ATENDE'} ao resultado mínimo de 1,00).",
        bold=True, center=True,
    )
    add_heading(
        doc,
        'FÓRMULA PARA FINS DE ATENDIMENTO AO ITEM "D.2" DA ALÍNEA "D" DO SUBITEM 11.1 '
        "DO ITEM 11 DO ANEXO VII-A DA IN SEGES/MP Nº 05/2017",
    )
    revenue = float(parameters.get("gross_revenue") or 0)
    variation = ((revenue - total) / revenue * 100) if revenue else 0
    add_body(
        doc,
        "Cálculo demonstrativo da variação percentual do valor total dos contratos firmados "
        "em relação à receita bruta discriminada na Demonstração do Resultado do Exercício (DRE).",
    )
    add_body(
        doc,
        f"({brl(revenue)} – {brl(total)}) × 100 ÷ {brl(revenue)} = {variation:.2f}%.",
        bold=True, center=True,
    )
    add_heading(doc, "Justificativa da variação superior a 10%", level=2)
    justification = parameters.get("justification_text") or (
        f"A divergência observada entre os valores apresentados na Demonstração do Resultado do "
        f"Exercício encerrada em 31 de dezembro de {parameters.get('reference_year')} e a relação "
        "de contratos firmados decorre da diferença nos critérios e nos períodos de reconhecimento "
        "das receitas. A DRE contempla as receitas efetivamente reconhecidas no exercício, enquanto "
        "a relação de contratos inclui todos os instrumentos vigentes, cujos faturamentos se "
        "distribuem ao longo de exercícios presentes e futuros."
    )
    add_body(doc, justification)
    add_body(doc, f"Brasília/DF, {today_brt().strftime('%d/%m/%Y')}.", center=True)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(3)
    add_signature_text(
        doc,
        "ENGEMIL ENGENHARIA, EMPREEENDIMENTOS, MANUTENÇÃO E INSTALAÇÕES LTDA",
        bold=True,
        border=True,
    )
    add_signature_text(doc, parameters.get("signatory_name") or "", bold=True)
    registration = " ".join(filter(None, [
        parameters.get("signatory_registration"),
        f"CPF: {parameters.get('signatory_cpf')}" if parameters.get("signatory_cpf") else None,
    ]))
    add_signature_text(doc, registration)
    add_signature_text(doc, parameters.get("signatory_title") or "")
    output = BytesIO()
    doc.save(output)
    return output.getvalue()
