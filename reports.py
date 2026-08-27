from __future__ import annotations

from datetime import date, datetime
from html import escape
from io import BytesIO
from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.utils import ImageReader
from reportlab.pdfbase.pdfmetrics import stringWidth
from reportlab.pdfgen import canvas as pdf_canvas

from contract_utils import extract_agency_acronym
from reportlab.platypus import Paragraph

BASE_DIR = Path(__file__).resolve().parent
DEFAULT_TEMPLATE = BASE_DIR / "templates" / "MODELO.docx"
BURGUNDY = "5A1235"
BURGUNDY_LIGHT = "F3E8ED"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "667085"
TABLE_INDENT_DXA = 100
CELL_MARGIN_DXA = 100


def brl(value) -> str:
    return (
        f"R$ {float(value or 0):,.2f}"
        .replace(",", "X")
        .replace(".", ",")
        .replace("X", ".")
    )


def percent(value, decimals=2) -> str:
    return (
        f"{float(value or 0):,.{decimals}f}%"
        .replace(",", "X")
        .replace(".", ",")
        .replace("X", ".")
    )


def short_date(value) -> str:
    if not value:
        return ""
    try:
        return date.fromisoformat(str(value)[:10]).strftime("%d/%m/%Y")
    except ValueError:
        return str(value)


def short_datetime(value) -> str:
    if not value:
        return ""
    try:
        return datetime.fromisoformat(str(value)).strftime("%d/%m/%Y %H:%M")
    except ValueError:
        return str(value)


def _meaningful(value) -> bool:
    return value is not None and str(value).strip() not in {"", "None", "nan", "NaT"}


def _set_run_font(run, *, bold=False, color="000000", size=11):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def _set_cell_fill(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), color)


def _set_cell_margins(cell, top=CELL_MARGIN_DXA, start=CELL_MARGIN_DXA,
                      bottom=CELL_MARGIN_DXA, end=CELL_MARGIN_DXA):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in (
        ("top", top), ("start", start), ("bottom", bottom), ("end", end),
    ):
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = tr_pr.find(qn("w:tblHeader"))
    if header is None:
        header = OxmlElement("w:tblHeader")
        tr_pr.append(header)
    header.set(qn("w:val"), "true")


def _prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def _set_table_geometry(table, widths_dxa):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table_pr = table._tbl.tblPr
    layout = table_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        table_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    width = table_pr.find(qn("w:tblW"))
    if width is None:
        width = OxmlElement("w:tblW")
        table_pr.append(width)
    width.set(qn("w:w"), str(sum(widths_dxa)))
    width.set(qn("w:type"), "dxa")
    indent = table_pr.find(qn("w:tblInd"))
    if indent is None:
        indent = OxmlElement("w:tblInd")
        table_pr.append(indent)
    indent.set(qn("w:w"), str(TABLE_INDENT_DXA))
    indent.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width_value in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width_value))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            value = widths_dxa[min(index, len(widths_dxa) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(value))
            tc_w.set(qn("w:type"), "dxa")
            _set_cell_margins(cell)


def _content_width_dxa(doc) -> int:
    section = doc.sections[0]
    # python-docx expõe a geometria em EMU; o OOXML das tabelas usa DXA (twips).
    return int(
        (section.page_width - section.left_margin - section.right_margin) / 635
    )


def _scaled_widths(doc, proportions):
    available = _content_width_dxa(doc) - TABLE_INDENT_DXA
    total = sum(proportions)
    widths = [round(available * value / total) for value in proportions]
    widths[-1] += available - sum(widths)
    return widths


def _append_word_field(paragraph, instruction: str, placeholder: str = "1"):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction_node = OxmlElement("w:instrText")
    instruction_node.set(qn("xml:space"), "preserve")
    instruction_node.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    result = OxmlElement("w:t")
    result.text = placeholder
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction_node, separate, result, end])
    _set_run_font(run, size=9)


def _rebuild_document_header(doc):
    logo_path = BASE_DIR / "assets" / "logo_engemil.png"

    def build_header_content(header):
        brand = header.add_paragraph()
        brand.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if logo_path.exists():
            brand.add_run().add_picture(str(logo_path), width=Mm(30))
        brand.paragraph_format.space_before = Pt(0)
        brand.paragraph_format.space_after = Pt(0)
        brand.paragraph_format.line_spacing = 1.0

        pagination = header.add_paragraph()
        pagination.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pagination.paragraph_format.space_before = Pt(0)
        pagination.paragraph_format.space_after = Pt(2)
        pagination.paragraph_format.line_spacing = 1.0
        _set_run_font(pagination.add_run("DECLARAÇÃO · Página "), bold=True, size=9)
        _append_word_field(pagination, "PAGE")
        _set_run_font(pagination.add_run(" de "), size=9)
        _append_word_field(pagination, "NUMPAGES")

        paragraph_borders = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "2")
        bottom.set(qn("w:color"), "000000")
        paragraph_borders.append(bottom)
        pagination._p.get_or_add_pPr().append(paragraph_borders)

    for section in doc.sections:
        section.different_first_page_header_footer = True

        # O cabeçalho "regular" (páginas 2 em diante) usava só um parágrafo
        # vazio, herdado de um desenho pensado para cartas de uma página só
        # (onde a diferença nunca aparecia). Em documentos de várias
        # páginas — como a Ficha Contratual — isso deixava a logo e a
        # identidade visual sumirem a partir da segunda página. Agora as
        # páginas seguintes recebem o mesmo cabeçalho da primeira.
        regular_header = section.header
        regular_header.is_linked_to_previous = False
        for child in list(regular_header._element):
            regular_header._element.remove(child)
        build_header_content(regular_header)

        first_header = section.first_page_header
        first_header.is_linked_to_previous = False
        for child in list(first_header._element):
            first_header._element.remove(child)
        build_header_content(first_header)

        section_properties = section._sectPr
        footer_references = section_properties.findall(qn("w:footerReference"))
        default_footer = next(
            (
                reference
                for reference in footer_references
                if reference.get(qn("w:type")) == "default"
            ),
            None,
        )
        if default_footer is not None:
            for reference in footer_references:
                if reference.get(qn("w:type")) == "first":
                    section_properties.remove(reference)
            first_footer = OxmlElement("w:footerReference")
            first_footer.set(qn("w:type"), "first")
            first_footer.set(qn("r:id"), default_footer.get(qn("r:id")))
            section_properties.insert(
                section_properties.index(default_footer) + 1,
                first_footer,
            )


def _configure_document(doc):
    for section in doc.sections:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Mm(210)
        section.page_height = Mm(297)
    _rebuild_document_header(doc)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), "Calibri")
    normal._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    return doc


def _new_document(template_path=DEFAULT_TEMPLATE):
    return _configure_document(Document(str(template_path)))


def _add_title(doc, text, subtitle=""):
    paragraph = doc.paragraphs[0] if doc.paragraphs and not doc.paragraphs[0].text else doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.line_spacing = 1.0
    _set_run_font(paragraph.add_run(text), bold=True, color=BURGUNDY)
    if subtitle:
        subtitle_paragraph = doc.add_paragraph()
        subtitle_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_paragraph.paragraph_format.space_before = Pt(0)
        subtitle_paragraph.paragraph_format.space_after = Pt(8)
        subtitle_paragraph.paragraph_format.line_spacing = 1.0
        _set_run_font(subtitle_paragraph.add_run(subtitle), color=MID_GRAY)


def _add_heading(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(10)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.paragraph_format.keep_with_next = True
    _set_run_font(paragraph.add_run(text), bold=True, color=BURGUNDY)
    return paragraph


def _add_text(doc, text, *, justified=False, bold=False, center=False, color="000000"):
    if not _meaningful(text):
        return None
    paragraph = doc.add_paragraph()
    if center:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif justified:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.25
    _set_run_font(paragraph.add_run(str(text)), bold=bold, color=color)
    return paragraph


def _clear_cell(cell):
    paragraph = cell.paragraphs[0]
    paragraph.clear()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    return paragraph


def _fill_lines(cell, lines, *, align=WD_ALIGN_PARAGRAPH.LEFT):
    paragraph = _clear_cell(cell)
    paragraph.alignment = align
    filtered = [(label, value) for label, value in lines if _meaningful(value)]
    if not filtered:
        _set_run_font(paragraph.add_run("—"), color=MID_GRAY)
        return
    for index, (label, value) in enumerate(filtered):
        if index:
            paragraph.add_run().add_break()
        if label:
            _set_run_font(paragraph.add_run(f"{label}: "), bold=True)
        _set_run_font(paragraph.add_run(str(value)))


def _add_header_row(table, headers):
    row = table.rows[0]
    _set_repeat_header(row)
    for cell, header in zip(row.cells, headers):
        _set_cell_fill(cell, BURGUNDY)
        paragraph = _clear_cell(cell)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_run_font(paragraph.add_run(header), bold=True, color="FFFFFF")


def _add_definition_table(doc, pairs):
    pairs = [(label, value) for label, value in pairs if _meaningful(value)]
    if not pairs:
        return None
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for label, value in pairs:
        row = table.add_row()
        _prevent_row_split(row)
        _set_cell_fill(row.cells[0], BURGUNDY_LIGHT)
        _fill_lines(row.cells[0], [("", label)])
        row.cells[0].paragraphs[0].runs[0].font.bold = True
        _fill_lines(row.cells[1], [("", value)])
    _set_table_geometry(table, _scaled_widths(doc, [1.35, 5.15]))
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def _write_cell_text(cell, value, *, align=WD_ALIGN_PARAGRAPH.LEFT):
    paragraph = _clear_cell(cell)
    paragraph.alignment = align
    for index, line in enumerate(str(value or "—").splitlines()):
        if index:
            paragraph.add_run().add_break()
        if ": " in line:
            label, detail = line.split(": ", 1)
            _set_run_font(paragraph.add_run(f"{label}: "), bold=True)
            _set_run_font(paragraph.add_run(detail))
        else:
            _set_run_font(paragraph.add_run(line))


def _add_records_table(doc, headers, rows, proportions):
    if not rows:
        return None
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    _add_header_row(table, headers)
    for values in rows:
        row = table.add_row()
        _prevent_row_split(row)
        for cell, value in zip(row.cells, values):
            _write_cell_text(cell, value)
    _set_table_geometry(table, _scaled_widths(doc, proportions))
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def _replace_header_label(data: bytes, label: str) -> bytes:
    source_buffer = BytesIO(data)
    target_buffer = BytesIO()
    with ZipFile(source_buffer) as source, ZipFile(target_buffer, "w", ZIP_DEFLATED) as target:
        for item in source.infolist():
            content = source.read(item.filename)
            if item.filename.startswith("word/header") and item.filename.endswith(".xml"):
                content = content.replace(
                    "DECLARAÇÃO".encode("utf-8"),
                    label.encode("utf-8"),
                )
            target.writestr(item, content)
    return target_buffer.getvalue()


def _save_bytes(doc, header_label=""):
    _configure_document(doc)
    output = BytesIO()
    doc.save(output)
    data = output.getvalue()
    return _replace_header_label(data, header_label) if header_label else data


def _pdf_text(value) -> str:
    """Normaliza texto para as fontes padrão do PDF sem perder acentuação comum."""
    return str(value or "").replace("\n", " ").replace("\r", " ").strip()


def _fit_pdf_text(value, width, font_name="Helvetica", font_size=4.0, min_size=2.8):
    """Ajusta texto a uma célula estreita, reduzindo a fonte antes de abreviar."""
    text = _pdf_text(value)
    if not text:
        return "", font_size
    available = max(float(width) - 3.0, 1.0)
    size = float(font_size)
    while size > min_size and stringWidth(text, font_name, size) > available:
        size = round(size - 0.15, 2)
    if stringWidth(text, font_name, size) <= available:
        return text, size
    suffix = "..."
    while text and stringWidth(text + suffix, font_name, size) > available:
        text = text[:-1]
    return text.rstrip() + suffix, size


def _draw_pdf_cell_text(
    surface,
    value,
    x,
    y,
    width,
    height,
    *,
    align="left",
    bold=False,
    font_size=4.0,
    color=colors.black,
):
    font_name = "Helvetica-Bold" if bold else "Helvetica"
    text, size = _fit_pdf_text(value, width, font_name, font_size)
    surface.setFont(font_name, size)
    surface.setFillColor(color)
    baseline = y + max((height - size) / 2.0, 1.2)
    if align == "right":
        surface.drawRightString(x + width - 1.5, baseline, text)
    elif align == "center":
        surface.drawCentredString(x + width / 2.0, baseline, text)
    else:
        surface.drawString(x + 1.5, baseline, text)


def _tint_hex(hex_color, factor):
    """Clareia uma cor em direção ao branco. factor=0 mantém a cor original,
    factor=1 resulta em branco. Usado para gerar paletas de um único matiz
    (mesma família de cor, tons diferentes) que continuam distinguíveis
    quando impressas em preto e branco — ao contrário de paletas com matizes
    diferentes (azul/verde/vermelho), que costumam colapsar em cinzas muito
    parecidos na conversão para escala de cinza."""
    factor = max(0.0, min(1.0, factor))
    r = int(hex_color[0:2], 16)
    g = int(hex_color[2:4], 16)
    b = int(hex_color[4:6], 16)
    r = int(r + (255 - r) * factor)
    g = int(g + (255 - g) * factor)
    b = int(b + (255 - b) * factor)
    return f"#{r:02x}{g:02x}{b:02x}"


def _readable_text_color(hex_color):
    """Decide preto ou branco para o texto sobre uma cor de fundo, pela
    luminância — garante contraste legível tanto em tela quanto impresso,
    inclusive em preto e branco (onde a luminância também determina o
    tom de cinza resultante)."""
    r = int(hex_color[1:3], 16)
    g = int(hex_color[3:5], 16)
    b = int(hex_color[5:7], 16)
    luminance = (0.299 * r + 0.587 * g + 0.114 * b) / 255
    return colors.white if luminance < 0.6 else colors.black


def build_contract_overview_summary(contracts, total_remaining_value=None):
    """Agrega os dados de uma lista de contratos (dicts no formato de
    `load_contracts()`) para alimentar a página de visão geral do PDF do
    Backlog. Mantém a lógica de agregação fora de reports.py para que este
    módulo só cuide de desenhar."""
    today = date.today()
    total_contracts = len(contracts)
    total_current_value = sum(float(c.get("current_value") or 0) for c in contracts)
    expiring_90_days = 0
    expiring_soon: list[dict] = []
    for contract in contracts:
        end_date = contract.get("end_date")
        if not end_date:
            continue
        try:
            days_left = (date.fromisoformat(str(end_date)[:10]) - today).days
        except ValueError:
            continue
        if 0 <= days_left <= 90:
            expiring_90_days += 1
            expiring_soon.append({
                "cost_center": contract.get("cost_center") or "—",
                "client": contract.get("client") or "—",
                "end_date": end_date,
                "days_left": days_left,
            })
    expiring_soon.sort(key=lambda item: item["days_left"])
    by_status: dict[str, int] = {}
    by_category_value: dict[str, float] = {}
    by_cost_center_value: dict[str, float] = {}
    by_instrument: dict[str, int] = {}
    for contract in contracts:
        status = contract.get("status") or "NÃO DEFINIDO"
        by_status[status] = by_status.get(status, 0) + 1
        category = contract.get("category") or "Não categorizado"
        by_category_value[category] = by_category_value.get(category, 0.0) + float(
            contract.get("current_value") or 0
        )
        cost_center = contract.get("cost_center") or "—"
        by_cost_center_value[cost_center] = by_cost_center_value.get(cost_center, 0.0) + float(
            contract.get("current_value") or 0
        )
    top_cost_centers = sorted(by_cost_center_value.items(), key=lambda kv: kv[1], reverse=True)[:8]
    top_categories = sorted(by_category_value.items(), key=lambda kv: kv[1], reverse=True)[:6]
    top_contracts = sorted(
        (
            {
                "cost_center": c.get("cost_center") or "—",
                "client": c.get("client") or "—",
                "current_value": float(c.get("current_value") or 0),
            }
            for c in contracts
        ),
        key=lambda item: item["current_value"], reverse=True,
    )[:8]
    average_value = (total_current_value / total_contracts) if total_contracts else 0.0
    return {
        "total_contracts": total_contracts,
        "total_current_value": total_current_value,
        "average_contract_value": average_value,
        "total_remaining_value": (
            float(total_remaining_value) if total_remaining_value is not None else None
        ),
        "expiring_90_days": expiring_90_days,
        "expiring_soon": expiring_soon[:8],
        "by_status": by_status,
        "top_cost_centers": top_cost_centers,
        "top_categories": top_categories,
        "top_contracts": top_contracts,
    }


def draw_contract_overview_page(surface, context, summary):
    """Desenha, em UMA página A4 horizontal, um dashboard da carteira de
    contratos: indicadores, maiores contratos, vencimentos próximos e
    distribuição por centro de custo/categoria/status. Pensado para
    impressão em preto e branco — usa uma única família de cor (variações
    de tom da cor institucional) em vez de matizes diferentes, porque tons
    diferentes de uma mesma cor continuam distinguíveis quando impressos
    sem cor; matizes diferentes costumam colapsar em cinzas parecidos. Todo
    texto sobre uma cor de fundo escolhe preto ou branco pela luminância do
    fundo, para nunca ficar ilegível."""
    page_width, page_height = landscape(A4)
    surface.setPageSize((page_width, page_height))
    margin_x = 36.0
    brand_logo = context["brand_logo"]
    reference_date = context["reference_date"]

    # ---- Cabeçalho — mesmas constantes de posicionamento da página 1, para
    # a logo nunca encostar na tarja do título (bug corrigido nesta versão:
    # a página 2 usava números próprios que deixavam só 3pt de folga real,
    # cortando a palavra ENGEMIL por baixo da tarja). ----
    cursor_y = page_height - 26.0
    logo_width = 66.0
    logo_height = 0.0
    if brand_logo.exists():
        try:
            logo_reader = ImageReader(str(brand_logo))
            logo_source_width, logo_source_height = logo_reader.getSize()
            logo_height = logo_width * logo_source_height / logo_source_width
            surface.drawImage(
                logo_reader, margin_x, cursor_y - logo_height, width=logo_width,
                height=logo_height, preserveAspectRatio=False, anchor="sw", mask="auto",
            )
            surface.setFillColor(colors.HexColor("#444444"))
            surface.setFont("Helvetica", 3.15)
            surface.drawCentredString(
                margin_x + logo_width / 2, cursor_y - logo_height - 4.8,
                "Engenharia, Empreendimentos, Manutenção e Instalações Ltda",
            )
        except Exception:
            pass
    cursor_y -= logo_height + 12.0

    bar_height = 17.0
    surface.setFillColor(colors.HexColor("#d9d9d9"))
    surface.rect(margin_x, cursor_y - bar_height, page_width - 2 * margin_x, bar_height, fill=1, stroke=0)
    surface.setFillColor(colors.black)
    surface.setFont("Helvetica-Bold", 12)
    surface.drawCentredString(page_width / 2, cursor_y - bar_height + 4.5, "VISÃO GERAL DA CARTEIRA DE CONTRATOS")
    cursor_y -= bar_height + 14.0

    surface.setFont("Helvetica", 8)
    surface.drawCentredString(
        page_width / 2, cursor_y,
        f"Posição em {reference_date.strftime('%d/%m/%Y')} — "
        f"{summary['total_contracts']} contrato(s) no filtro selecionado",
    )
    surface.setFont("Helvetica", 6.6)
    surface.drawRightString(
        page_width - margin_x, cursor_y,
        f"Página {context['page_number']} de {context['page_count']}",
    )
    cursor_y -= 22.0

    # ---- Cartões de indicadores ----
    kpi_height = 44.0
    kpi_values = [
        ("CONTRATOS NO FILTRO", str(summary["total_contracts"]), None),
        ("VALOR VIGENTE TOTAL", brl(summary["total_current_value"]), None),
        (
            "REMANESCENTE TOTAL",
            brl(summary["total_remaining_value"]) if summary["total_remaining_value"] is not None else "—",
            None,
        ),
        ("TICKET MÉDIO POR CONTRATO", brl(summary["average_contract_value"]), None),
        ("VENCENDO EM 90 DIAS", str(summary["expiring_90_days"]), "alerta" if summary["expiring_90_days"] else None),
    ]
    kpi_gap = 8.0
    kpi_width = (page_width - 2 * margin_x - 4 * kpi_gap) / 5
    x = margin_x
    for label, value, flag in kpi_values:
        surface.setStrokeColor(colors.black)
        surface.setLineWidth(0.7)
        surface.setFillColor(colors.HexColor("#f7f7f7"))
        surface.rect(x, cursor_y - kpi_height, kpi_width, kpi_height, fill=1, stroke=1)
        surface.setFillColor(colors.HexColor("#444444"))
        surface.setFont("Helvetica-Bold", 6.3)
        _draw_pdf_cell_text(surface, label, x, cursor_y - 13, kpi_width, 10, align="center", font_size=6.3)
        surface.setFillColor(colors.HexColor(f"#{BURGUNDY}") if flag else colors.black)
        surface.setFont("Helvetica-Bold", 14)
        surface.drawCentredString(x + kpi_width / 2, cursor_y - 33, value)
        x += kpi_width + kpi_gap
    cursor_y -= kpi_height + 18.0

    # ---- Gráficos de barras horizontais (mesmo matiz, tons diferentes) ----
    chart_gap = 22.0
    chart_width = (page_width - 2 * margin_x - chart_gap) / 2
    chart_panel_height = 148.0
    charts_top = cursor_y

    def draw_bar_chart(x0, title, entries, value_formatter):
        surface.setFillColor(colors.black)
        surface.setFont("Helvetica-Bold", 8.5)
        surface.drawString(x0, charts_top, title)
        available_height = chart_panel_height - 16
        if not entries:
            surface.setFont("Helvetica", 7)
            surface.setFillColor(colors.HexColor("#667085"))
            surface.drawString(x0, charts_top - 20, "Sem dados para este filtro.")
            return
        row_h = min(20.0, available_height / len(entries))
        max_value = max(value for _, value in entries) or 1
        label_width = chart_width * 0.34
        bar_area_width = chart_width - label_width - 55
        y = charts_top - 14
        for index, (label, value) in enumerate(entries):
            y -= row_h
            tone = _tint_hex(BURGUNDY, index / max(1, len(entries) - 1) * 0.78)
            bar_w = max(2.0, bar_area_width * (value / max_value))
            surface.setFont("Helvetica", 6.6)
            surface.setFillColor(colors.black)
            _draw_pdf_cell_text(
                surface, str(label), x0, y, label_width, row_h - 2,
                align="left", font_size=6.6,
            )
            surface.setStrokeColor(colors.black)
            surface.setLineWidth(0.4)
            surface.setFillColor(colors.HexColor(tone))
            bar_y = y + (row_h - 2 - 9) / 2 + 1
            surface.rect(x0 + label_width, bar_y, bar_w, 9, fill=1, stroke=1)
            surface.setFont("Helvetica-Bold", 6.6)
            surface.setFillColor(colors.black)
            surface.drawString(x0 + label_width + bar_w + 4, bar_y + 2.2, value_formatter(value))

    draw_bar_chart(
        margin_x,
        "VALOR VIGENTE POR CENTRO DE CUSTO (maiores 8)",
        summary["top_cost_centers"],
        brl,
    )
    draw_bar_chart(
        margin_x + chart_width + chart_gap,
        "VALOR VIGENTE POR CATEGORIA",
        summary["top_categories"],
        brl,
    )
    cursor_y -= chart_panel_height + 16.0

    # ---- Tabelas: maiores contratos e vencimentos próximos ----
    table_gap = 22.0
    table_width = (page_width - 2 * margin_x - table_gap) / 2
    table_panel_height = 132.0
    tables_top = cursor_y

    def draw_mini_table(x0, title, headers, col_widths, aligns, data_rows):
        surface.setFillColor(colors.black)
        surface.setFont("Helvetica-Bold", 8.5)
        surface.drawString(x0, tables_top, title)
        header_y = tables_top - 16
        header_h = 13.0
        surface.setFillColor(colors.HexColor(f"#{BURGUNDY}"))
        surface.rect(x0, header_y - header_h, table_width, header_h, fill=1, stroke=0)
        header_text_color = _readable_text_color(f"#{BURGUNDY}")
        surface.setFillColor(header_text_color)
        surface.setFont("Helvetica-Bold", 6.6)
        cx = x0
        for header, width, align in zip(headers, col_widths, aligns):
            _draw_pdf_cell_text(
                surface, header, cx, header_y - header_h + 2.5, width, 9,
                align=align, font_size=6.6, color=header_text_color, bold=True,
            )
            cx += width
        row_top = header_y - header_h
        available_rows_height = table_panel_height - 16 - header_h
        if not data_rows:
            surface.setFillColor(colors.HexColor("#667085"))
            surface.setFont("Helvetica", 7)
            surface.drawString(x0, row_top - 14, "Sem registros para este filtro.")
            return
        row_h = min(14.0, available_rows_height / len(data_rows))
        y = row_top
        for index, row_values in enumerate(data_rows):
            y -= row_h
            if index % 2 == 1:
                surface.setFillColor(colors.HexColor("#f3e8ed"))
                surface.rect(x0, y, table_width, row_h, fill=1, stroke=0)
            surface.setFillColor(colors.black)
            surface.setFont("Helvetica", 6.4)
            cx = x0
            for value, width, align in zip(row_values, col_widths, aligns):
                _draw_pdf_cell_text(surface, str(value), cx, y + 2, width, row_h - 2, align=align, font_size=6.4)
                cx += width
            surface.setStrokeColor(colors.HexColor("#e1e3e7"))
            surface.setLineWidth(0.3)
            surface.line(x0, y, x0 + table_width, y)

    top_contracts_cols = [table_width * 0.20, table_width * 0.52, table_width * 0.28]
    draw_mini_table(
        margin_x, "MAIORES CONTRATOS (valor vigente)",
        ["CENTRO DE CUSTO", "CONTRATANTE", "VALOR ATUAL"],
        top_contracts_cols, ["left", "left", "right"],
        [
            (c["cost_center"], c["client"], brl(c["current_value"]))
            for c in summary["top_contracts"]
        ],
    )
    expiring_cols = [table_width * 0.18, table_width * 0.47, table_width * 0.17, table_width * 0.18]
    draw_mini_table(
        margin_x + table_width + table_gap, "VENCENDO NOS PRÓXIMOS 90 DIAS",
        ["CENTRO DE CUSTO", "CONTRATANTE", "FIM", "DIAS"],
        expiring_cols, ["left", "left", "center", "center"],
        [
            (
                e["cost_center"], e["client"],
                date.fromisoformat(str(e["end_date"])[:10]).strftime("%d/%m/%Y"),
                e["days_left"],
            )
            for e in summary["expiring_soon"]
        ],
    )
    cursor_y -= table_panel_height + 14.0

    # ---- Tarja de status (contagem) ----
    status_top = cursor_y
    if summary["by_status"]:
        surface.setFont("Helvetica-Bold", 7.5)
        surface.setFillColor(colors.black)
        surface.drawString(margin_x, status_top, "CONTRATOS POR STATUS:")
        chip_x = margin_x + 118
        ordered_status = sorted(summary["by_status"].items(), key=lambda kv: kv[1], reverse=True)
        for index, (status_label, count) in enumerate(ordered_status):
            tone = _tint_hex(BURGUNDY, index / max(1, len(ordered_status) - 1) * 0.7)
            text = f"{status_label}: {count}"
            text_width = stringWidth(text, "Helvetica-Bold", 6.8)
            chip_width = text_width + 14
            if chip_x + chip_width > page_width - margin_x:
                break
            surface.setFillColor(colors.HexColor(tone))
            surface.setStrokeColor(colors.black)
            surface.setLineWidth(0.4)
            surface.roundRect(chip_x, status_top - 12, chip_width, 14, 3, fill=1, stroke=1)
            surface.setFillColor(_readable_text_color(tone))
            surface.setFont("Helvetica-Bold", 6.8)
            surface.drawCentredString(chip_x + chip_width / 2, status_top - 8, text)
            chip_x += chip_width + 8

    surface.setStrokeColor(colors.HexColor("#a6a6a6"))
    surface.setLineWidth(0.35)
    surface.line(margin_x, 34, page_width - margin_x, 34)
    surface.setFillColor(colors.HexColor("#555555"))
    surface.setFont("Helvetica", 6.5)
    surface.drawCentredString(
        page_width / 2, 22,
        "ENGEMIL - Engenharia, Empreendimentos, Manutenção e Instalações Ltda. "
        "- CNPJ 04.768.702/0001-70 - gestao.contratos@engemileng.com",
    )
    surface.drawRightString(page_width - margin_x, 22, reference_date.strftime("%d/%m/%Y"))


def generate_backlog_pdf(
    rows,
    reference_date=None,
    logo_path=None,
    signatory=None,
    sort_label=None,
    report_title="BACKLOG ENGEMIL - POSIÇÃO DA CARTEIRA DE CONTRATOS",
    report_subtitle=None,
    pdf_subject="Posição da carteira de contratos",
    force_single_table_page=False,
    additional_page_renderer=None,
    overview_summary=None,
):
    """Gera o backlog pronto para envio no padrão compacto da planilha histórica.

    Quando `overview_summary` é informado (ver `build_contract_overview_summary`)
    e nenhum `additional_page_renderer` explícito foi passado, uma página extra
    em A4 horizontal é acrescentada com um dashboard de uma página só, pensado
    para impressão em preto e branco.
    """
    reference_date = reference_date or date.today()
    rows = [dict(row) for row in rows]
    signatory = dict(signatory or {})
    if overview_summary and additional_page_renderer is None:
        additional_page_renderer = lambda surface, context: draw_contract_overview_page(
            surface, context, overview_summary
        )
    output = BytesIO()
    page_width, page_height = A4
    surface = pdf_canvas.Canvas(output, pagesize=A4, pageCompression=1)
    metadata_title = (
        "Backlog ENGEMIL"
        if report_title == "BACKLOG ENGEMIL - POSIÇÃO DA CARTEIRA DE CONTRATOS"
        else _pdf_text(report_title)
    )
    surface.setTitle(metadata_title)
    surface.setAuthor("ENGEMIL Engenharia, Empreendimentos, Manutenção e Instalações Ltda.")
    surface.setSubject(_pdf_text(pdf_subject))

    margin_x = 27.5
    table_width = page_width - (2 * margin_x)
    column_widths = [20, 50, 151, 49, 35, 35, 60, 71, 68]
    column_widths[-1] += table_width - sum(column_widths)
    columns = [
        ("ITEM", "center"),
        ("CENTRO DE CUSTO", "center"),
        ("NOME DO CONTRATANTE", "left"),
        ("Nº / ANO DO CONTRATO", "center"),
        ("INÍCIO", "center"),
        ("FIM", "center"),
        ("VALOR ATUAL", "right"),
        ("INSTRUMENTO VIGENTE", "center"),
        ("REMANESCENTE TOTAL", "right"),
    ]
    header_height = 23.0
    has_signature = bool(signatory.get("name"))
    # A assinatura fica ancorada no rodapé, com uma área livre acima da linha
    # para inserção de assinatura eletrônica. O tamanho das linhas se adapta à
    # quantidade de contratos para manter a carteira corrente em uma página.
    signature_line_y = 70.0
    signature_table_limit = 125.0 if has_signature else 43.0
    total_height = 12.0
    table_top = page_height - 121.0
    table_bottom = signature_table_limit
    available_rows_height = max(
        1.0,
        table_top - table_bottom - header_height - total_height,
    )
    row_height = min(9.45, available_rows_height / max(1, len(rows)))
    # Ajuste automático: tenta sempre caber a carteira inteira em uma única
    # página, mas sem passar do limite de legibilidade para impressão. Só
    # quando a quantidade de contratos obrigaria a linhas menores que esse
    # limite é que o relatório volta a paginar normalmente.
    PRINT_LEGIBLE_ROW_HEIGHT = 3.6
    natural_single_page_height = available_rows_height / max(1, len(rows))
    auto_fits_single_page = natural_single_page_height >= PRINT_LEGIBLE_ROW_HEIGHT
    if force_single_table_page:
        row_height = max(2.8, row_height)
        capacity = max(1, len(rows))
    elif auto_fits_single_page:
        row_height = min(9.45, max(PRINT_LEGIBLE_ROW_HEIGHT, natural_single_page_height))
        capacity = max(1, len(rows))
    else:
        row_height = max(6.45, row_height)
        capacity = max(
            1,
            int(
                (
                    table_top - table_bottom - header_height
                    - total_height
                )
                // row_height
            ),
        )
    page_chunks = [
        rows[index:index + capacity]
        for index in range(0, len(rows), capacity)
    ] or [[]]
    total_current = sum(float(row.get("Valor atual") or 0) for row in rows)
    total_remaining = sum(float(row.get("Remanescente total") or 0) for row in rows)
    page_count = len(page_chunks)
    displayed_page_count = page_count + (1 if additional_page_renderer else 0)
    brand_logo = Path(logo_path) if logo_path else BASE_DIR / "assets" / "logo_engemil.png"

    def draw_brand(page_number):
        if brand_logo.exists():
            try:
                logo_reader = ImageReader(str(brand_logo))
                logo_source_width, logo_source_height = logo_reader.getSize()
                logo_width = 66.0
                logo_height = logo_width * logo_source_height / logo_source_width
                surface.drawImage(
                    logo_reader,
                    margin_x,
                    page_height - 47.0,
                    width=logo_width,
                    height=logo_height,
                    preserveAspectRatio=False,
                    anchor="sw",
                    mask="auto",
                )
                # O arquivo escuro original não traz a razão social sob o
                # logotipo. Ela é acrescentada aqui para que o cabeçalho exiba
                # a identidade visual completa, sem corte.
                surface.setFillColor(colors.HexColor("#444444"))
                surface.setFont("Helvetica", 3.15)
                surface.drawCentredString(
                    margin_x + logo_width / 2,
                    page_height - 51.8,
                    "Engenharia, Empreendimentos, Manutenção e Instalações Ltda",
                )
            except Exception:
                pass
        surface.setFillColor(colors.HexColor("#d9d9d9"))
        surface.rect(margin_x, page_height - 72, table_width, 17, fill=1, stroke=0)
        surface.setFillColor(colors.black)
        fitted_title, fitted_title_size = _fit_pdf_text(
            report_title,
            table_width - 12,
            "Helvetica-Bold",
            7.2,
            5.2,
        )
        surface.setFont("Helvetica-Bold", fitted_title_size)
        surface.drawCentredString(page_width / 2, page_height - 66.5, fitted_title)
        surface.setFont("Helvetica", 4.8)
        declaration = (
            "ENGEMIL - Engenharia, Empreendimentos, Manutenção e Instalações Ltda. "
            "- CNPJ 04.768.702/0001-70"
        )
        surface.drawCentredString(page_width / 2, page_height - 82, declaration)
        surface.drawCentredString(
            page_width / 2,
            page_height - 89,
            _pdf_text(report_subtitle) if report_subtitle else
            f"Contratos do filtro selecionado - posição em {reference_date.strftime('%d/%m/%Y')}",
        )
        if sort_label:
            surface.setFont("Helvetica", 4.1)
            surface.drawCentredString(
                page_width / 2,
                page_height - 96,
                f"Ordenação: {_pdf_text(sort_label)}",
            )
        surface.setFont("Helvetica", 4.2)
        surface.drawRightString(
            page_width - margin_x,
            page_height - 102,
            f"Página {page_number} de {displayed_page_count}",
        )

    def draw_header(y_top):
        surface.setStrokeColor(colors.black)
        surface.setLineWidth(0.35)
        x = margin_x
        y = y_top - header_height
        for index, ((label, _), width) in enumerate(zip(columns, column_widths)):
            surface.setFillColor(colors.HexColor("#d9d9d9"))
            surface.rect(x, y, width, header_height, fill=1, stroke=1)
            if index in {4, 5}:
                surface.setFont("Helvetica-Bold", 4.4)
                surface.setFillColor(colors.black)
                surface.drawCentredString(x + width / 2, y + 5.1, label)
            else:
                _draw_pdf_cell_text(
                    surface,
                    label,
                    x,
                    y,
                    width,
                    header_height,
                    align="center",
                    bold=True,
                    font_size=4.35,
                )
            x += width
        surface.setFont("Helvetica-Bold", 4.3)
        surface.drawCentredString(
            margin_x + sum(column_widths[:4]) + (column_widths[4] + column_widths[5]) / 2,
            y + 15.4,
            "VIGÊNCIA",
        )
        return y

    def draw_footer():
        surface.setStrokeColor(colors.HexColor("#a6a6a6"))
        surface.setLineWidth(0.35)
        surface.line(margin_x, 35.5, page_width - margin_x, 35.5)
        surface.setFillColor(colors.HexColor("#555555"))
        surface.setFont("Helvetica", 3.8)
        surface.drawCentredString(
            page_width / 2,
            29,
            'CRS 503, BLOCO "B", LOJA 05, PARTE SUPERIOR "A" - BRASÍLIA/DF - CEP 70.331-520',
        )
        surface.drawCentredString(
            page_width / 2,
            23.5,
            "(61) 3248-2876 / 3248-5919 / 3248-3410",
        )
        surface.drawCentredString(
            page_width / 2,
            18,
            "gestao.contratos@engemileng.com - https://engemil.com.br/",
        )
        surface.drawRightString(
            page_width - margin_x,
            18,
            reference_date.strftime("%d/%m/%Y"),
        )

    for page_number, chunk in enumerate(page_chunks, start=1):
        draw_brand(page_number)
        y = draw_header(table_top)
        for offset, row in enumerate(chunk):
            y -= row_height
            surface.setFillColor(
                colors.HexColor("#f2f2f2") if offset % 2 == 0 else colors.white
            )
            surface.rect(margin_x, y, table_width, row_height, fill=1, stroke=0)
            values = [
                f"{int(row.get('Item') or 0):02d}",
                row.get("Centro de custo"),
                row.get("Contratante"),
                row.get("Contrato"),
                short_date(row.get("Início")),
                short_date(row.get("Fim")),
                brl(row.get("Valor atual")),
                row.get("Instrumento vigente"),
                brl(row.get("Remanescente total")),
            ]
            x = margin_x
            for value, width, (_, align) in zip(values, column_widths, columns):
                surface.setStrokeColor(colors.black)
                surface.setLineWidth(0.22)
                surface.rect(x, y, width, row_height, fill=0, stroke=1)
                _draw_pdf_cell_text(
                    surface,
                    value,
                    x,
                    y,
                    width,
                    row_height,
                    align=align,
                    font_size=min(4.05, max(2.6, row_height - 1.8)),
                )
                x += width
        if page_number == page_count:
            y -= total_height
            surface.setFillColor(colors.HexColor("#d9d9d9"))
            surface.rect(margin_x, y, table_width, total_height, fill=1, stroke=1)
            surface.setStrokeColor(colors.black)
            x_value = margin_x + sum(column_widths[:6])
            x_remaining = margin_x + sum(column_widths[:8])
            _draw_pdf_cell_text(
                surface,
                f"VALOR DOS CONTRATOS ({len(rows)} contratos)",
                margin_x,
                y,
                sum(column_widths[:6]),
                total_height,
                align="right",
                bold=True,
                font_size=4.4,
            )
            _draw_pdf_cell_text(
                surface,
                brl(total_current),
                x_value,
                y,
                column_widths[6],
                total_height,
                align="right",
                bold=True,
                font_size=4.2,
            )
            _draw_pdf_cell_text(
                surface,
                brl(total_remaining),
                x_remaining,
                y,
                column_widths[8],
                total_height,
                align="right",
                bold=True,
                font_size=4.2,
            )
            if has_signature:
                signature_center = page_width / 2
                signature_width = 235.0
                surface.setStrokeColor(colors.black)
                surface.setLineWidth(0.45)
                surface.line(
                    signature_center - signature_width / 2,
                    signature_line_y,
                    signature_center + signature_width / 2,
                    signature_line_y,
                )
                surface.setFillColor(colors.black)
                surface.setFont("Helvetica-Bold", 5.2)
                surface.drawCentredString(
                    signature_center,
                    signature_line_y - 8,
                    _pdf_text(signatory.get("name")),
                )
                title = _pdf_text(signatory.get("title"))
                if title:
                    surface.setFont("Helvetica", 4.5)
                    surface.drawCentredString(
                        signature_center,
                        signature_line_y - 15,
                        title,
                    )
                professional_data = " - ".join(
                    value for value in (
                        _pdf_text(signatory.get("registration")),
                        (
                            f"CPF {_pdf_text(signatory.get('cpf'))}"
                            if signatory.get("cpf") else ""
                        ),
                    )
                    if value
                )
                if professional_data:
                    surface.setFont("Helvetica", 4.2)
                    surface.drawCentredString(
                        signature_center,
                        signature_line_y - 22,
                        professional_data,
                    )
        draw_footer()
        surface.showPage()
    if additional_page_renderer:
        additional_page_renderer(
            surface,
            {
                "page_width": page_width,
                "page_height": page_height,
                "margin_x": margin_x,
                "table_width": table_width,
                "brand_logo": brand_logo,
                "reference_date": reference_date,
                "page_number": page_count + 1,
                "page_count": displayed_page_count,
                "report_title": report_title,
            },
        )
        surface.showPage()
    surface.save()
    return output.getvalue()


BID_PDF_COLUMN_CATALOG = {
    "process_number": ("PROCESSO", "center", 1.1, False),
    "uasg": ("UASG", "center", 0.55, False),
    "edital_number": ("EDITAL", "center", 0.75, False),
    "agency": ("ÓRGÃO", "center", 1.3, True),
    "uf": ("UF", "center", 0.4, False),
    "platform": ("PLATAFORMA", "center", 1.1, True),
    "modality": ("MODALIDADE", "center", 1.0, True),
    "scope": ("ESCOPO", "center", 0.9, False),
    "structure": ("ESTRUTURA", "center", 1.4, True),
    "object": ("OBJETO", "left", 2.9, True),
    "estimated_value": ("VALOR ESTIMADO", "right", 1.05, True),
    "our_bid_value": ("NOSSO LANCE", "right", 1.0, False),
    "our_discount_percent": ("DESCONTO", "right", 0.7, False),
    "our_ranking": ("CLASSIFICAÇÃO", "center", 0.8, False),
    "status": ("STATUS", "center", 1.2, False),
    "dispute_date": ("DATA DA DISPUTA", "center", 0.85, False),
    "dispute_time": ("HORÁRIO", "center", 0.5, False),
    "responsible_name": ("RESPONSÁVEL", "center", 1.1, False),
}


def _wrap_pdf_text_lines(text, width, font_name="Helvetica", font_size=7.0, max_lines=4):
    """Quebra um texto em várias linhas dentro da largura disponível, em vez
    de diminuir a fonte — usado nas colunas onde o conteúdo precisa
    continuar legível (ex.: Objeto, Estrutura). Respeita quebras de linha
    já embutidas no texto (ex.: "Individual\\n(item único)") como pontos de
    quebra "de propósito", além de quebrar automaticamente por largura
    dentro de cada trecho. Trunca com reticências além de `max_lines` para
    a linha não crescer sem limite."""
    raw = str(text or "").replace("\r", "").strip()
    if not raw:
        return [""]
    available = max(float(width) - 6.0, 10.0)
    lines = []
    for segment in raw.split("\n"):
        segment = segment.strip()
        if not segment:
            continue
        words = segment.split()
        current = ""
        for word in words:
            candidate = f"{current} {word}".strip()
            if stringWidth(candidate, font_name, font_size) > available and current:
                lines.append(current)
                current = word
            else:
                current = candidate
        if current:
            lines.append(current)
    if len(lines) > max_lines:
        lines = lines[:max_lines]
        last = lines[-1]
        while stringWidth(last + "…", font_name, font_size) > available and len(last) > 1:
            last = last[:-1]
        lines[-1] = last + "…"
    return lines or [""]


def generate_bid_processes_pdf(
    rows,
    column_keys,
    reference_date=None,
    logo_path=None,
    report_title="LICITAÇÕES VIGENTES ENGEMIL",
    report_subtitle=None,
):
    """Gera o relatório de licitações vigentes em A4 horizontal, ajustando
    automaticamente para uma página só quando possível (mesmo princípio do
    Backlog oficial — só pagina se isso for necessário), com as colunas
    escolhidas pelo usuário. `rows` é uma lista de dicts no formato de
    `bid_processes`; `column_keys` é a lista, na ordem desejada, de chaves
    presentes em BID_PDF_COLUMN_CATALOG.

    Colunas marcadas para quebra de linha (hoje, Objeto) nunca diminuem a
    fonte para caber — em vez disso, a linha da tabela cresce em altura
    conforme o texto precisa de mais linhas, mantendo tudo legível."""
    reference_date = reference_date or date.today()
    rows = [dict(row) for row in rows]
    columns = [
        (key, *BID_PDF_COLUMN_CATALOG[key])
        for key in column_keys if key in BID_PDF_COLUMN_CATALOG
    ] or [("process_number", *BID_PDF_COLUMN_CATALOG["process_number"])]

    output = BytesIO()
    page_width, page_height = landscape(A4)
    surface = pdf_canvas.Canvas(output, pagesize=landscape(A4), pageCompression=1)
    surface.setTitle("Licitações vigentes ENGEMIL")
    surface.setAuthor("ENGEMIL Engenharia, Empreendimentos, Manutenção e Instalações Ltda.")
    surface.setSubject(_pdf_text("Posição das licitações em andamento"))

    margin_x = 27.5
    table_width = page_width - (2 * margin_x)
    weight_total = sum(weight for _, _, _, weight, _ in columns)
    column_widths = [table_width * (weight / weight_total) for _, _, _, weight, _ in columns]

    brand_logo = Path(logo_path) if logo_path else BASE_DIR / "assets" / "logo_engemil.png"
    CELL_FONT_SIZE = 7.0
    LINE_HEIGHT = 8.6
    MIN_ROW_HEIGHT = 15.0
    header_height = 15.0
    table_top = page_height - 108.0
    table_bottom = 40.0

    def format_structure_for_pdf(label):
        """Quebra o resumo de estrutura num ponto de sentido (não só por
        largura), para ficar em duas linhas dentro da célula — ex.:
        'Individual (item único)' vira 'Individual' / '(item único)'."""
        text = str(label or "")
        if " · " in text:
            return text.replace(" · ", "\n", 1)
        if " (" in text:
            return text.replace(" (", "\n(", 1)
        return text

    def format_cell(key, row):
        value = row.get(key)
        if key == "agency":
            acronym = extract_agency_acronym(str(value or ""))
            return acronym or str(value or "—")
        if key == "structure":
            return format_structure_for_pdf(value)
        if key in ("estimated_value", "our_bid_value"):
            if isinstance(value, str):
                return value
            return brl(value) if value else "—"
        if key == "our_discount_percent":
            if value is None:
                return "—"
            if value < 0:
                return f"▲ {abs(value):.2f}%".replace(".", ",")
            return f"{value:.2f}%".replace(".", ",")
        if key == "dispute_date" and value:
            try:
                return date.fromisoformat(str(value)[:10]).strftime("%d/%m/%Y")
            except ValueError:
                return str(value)
        return str(value) if value not in (None, "") else "—"

    def row_height_for(row):
        needed_lines = 1
        for (key, _, _, width, wrap), col_width in zip(columns, column_widths):
            if not wrap:
                continue
            lines = _wrap_pdf_text_lines(format_cell(key, row), col_width, font_size=CELL_FONT_SIZE)
            needed_lines = max(needed_lines, len(lines))
        return max(MIN_ROW_HEIGHT, needed_lines * LINE_HEIGHT + 6.0)

    row_heights = [row_height_for(row) for row in rows]
    available_rows_height = max(1.0, table_top - table_bottom - header_height)

    page_chunks = []
    current_chunk = []
    current_height = 0.0
    for row, height in zip(rows, row_heights):
        if current_chunk and current_height + height > available_rows_height:
            page_chunks.append(current_chunk)
            current_chunk = []
            current_height = 0.0
        current_chunk.append((row, height))
        current_height += height
    if current_chunk or not rows:
        page_chunks.append(current_chunk)
    page_count = len(page_chunks)

    def draw_brand(page_number):
        if brand_logo.exists():
            try:
                logo_reader = ImageReader(str(brand_logo))
                source_w, source_h = logo_reader.getSize()
                logo_width = 66.0
                logo_height = logo_width * source_h / source_w
                surface.drawImage(
                    logo_reader, margin_x, page_height - 47.0, width=logo_width,
                    height=logo_height, preserveAspectRatio=False, anchor="sw", mask="auto",
                )
                surface.setFillColor(colors.HexColor("#444444"))
                surface.setFont("Helvetica", 3.15)
                surface.drawCentredString(
                    margin_x + logo_width / 2, page_height - 51.8,
                    "Engenharia, Empreendimentos, Manutenção e Instalações Ltda",
                )
            except Exception:
                pass
        surface.setFillColor(colors.HexColor("#d9d9d9"))
        surface.rect(margin_x, page_height - 72, table_width, 17, fill=1, stroke=0)
        surface.setFillColor(colors.black)
        surface.setFont("Helvetica-Bold", 11)
        surface.drawCentredString(page_width / 2, page_height - 66.5, report_title)
        surface.setFont("Helvetica", 6.5)
        surface.drawCentredString(
            page_width / 2, page_height - 82,
            "ENGEMIL - Engenharia, Empreendimentos, Manutenção e Instalações Ltda. "
            "- CNPJ 04.768.702/0001-70",
        )
        surface.drawCentredString(
            page_width / 2, page_height - 91,
            _pdf_text(report_subtitle) if report_subtitle else
            f"{len(rows)} licitação(ões) - posição em {reference_date.strftime('%d/%m/%Y')}",
        )
        surface.setFont("Helvetica", 6.0)
        surface.drawRightString(page_width - margin_x, page_height - 100, f"Página {page_number} de {page_count}")

    def draw_header(y_top):
        x = margin_x
        y = y_top - header_height
        surface.setStrokeColor(colors.black)
        surface.setLineWidth(0.35)
        for (key, label, align, _, _), width in zip(columns, column_widths):
            surface.setFillColor(colors.HexColor(f"#{BURGUNDY}"))
            surface.rect(x, y, width, header_height, fill=1, stroke=1)
            _draw_pdf_cell_text(
                surface, label, x, y, width, header_height, align="center",
                font_size=6.6, color=_readable_text_color(f"#{BURGUNDY}"), bold=True,
            )
            x += width
        return y

    for page_index, page_rows in enumerate(page_chunks):
        draw_brand(page_index + 1)
        y = draw_header(table_top)
        for row_index, (row, row_height) in enumerate(page_rows):
            y -= row_height
            if row_index % 2 == 1:
                surface.setFillColor(colors.HexColor("#f7f7f7"))
                surface.rect(margin_x, y, table_width, row_height, fill=1, stroke=0)
            x = margin_x
            for (key, _, align, _, wrap), width in zip(columns, column_widths):
                cell_text = format_cell(key, row)
                if wrap:
                    lines = _wrap_pdf_text_lines(cell_text, width, font_size=CELL_FONT_SIZE)
                    block_height = len(lines) * LINE_HEIGHT
                    top_gap = max((row_height - block_height) / 2.0, 2.0)
                    line_y = y + row_height - top_gap - LINE_HEIGHT + 2
                    surface.setFont("Helvetica", CELL_FONT_SIZE)
                    surface.setFillColor(colors.black)
                    for line in lines:
                        line_w = stringWidth(line, "Helvetica", CELL_FONT_SIZE)
                        if align == "center":
                            line_x = x + (width - line_w) / 2.0
                        elif align == "right":
                            line_x = x + width - line_w - 4
                        else:
                            line_x = x + 3
                        surface.drawString(line_x, line_y, line)
                        line_y -= LINE_HEIGHT
                else:
                    cell_color = colors.black
                    if key == "our_discount_percent" and row.get(key) is not None:
                        cell_color = (
                            colors.HexColor("#B5651D") if row[key] < 0 else colors.black
                        )
                    _draw_pdf_cell_text(
                        surface, cell_text, x, y, width, row_height,
                        align=align, font_size=CELL_FONT_SIZE, color=cell_color,
                    )
                x += width
            surface.setStrokeColor(colors.HexColor("#e1e3e7"))
            surface.setLineWidth(0.3)
            surface.line(margin_x, y, margin_x + table_width, y)
        surface.setStrokeColor(colors.HexColor("#a6a6a6"))
        surface.setLineWidth(0.35)
        surface.line(margin_x, 30, page_width - margin_x, 30)
        surface.setFillColor(colors.HexColor("#555555"))
        surface.setFont("Helvetica", 6.0)
        surface.drawCentredString(
            page_width / 2, 18,
            "gestao.contratos@engemileng.com - https://engemil.com.br/",
        )
        surface.drawRightString(page_width - margin_x, 18, reference_date.strftime("%d/%m/%Y"))
        surface.showPage()
    surface.save()
    return output.getvalue()


def generate_indices_pdf(rows, parameters, reference_date=None, logo_path=None):
    """Gera a declaração de índices em duas páginas A4 verticais.

    A primeira página usa a mesma tabela oficial do Backlog. A segunda concentra
    as fórmulas, a justificativa e a assinatura, evitando mudanças de orientação.
    """
    reference_date = reference_date or date.today()
    parameters = dict(parameters or {})
    report_title = "DECLARAÇÃO DE CONTRATOS E ÍNDICES ECONÔMICO-FINANCEIROS"

    def draw_details_page(surface, context):
        page_width = context["page_width"]
        page_height = context["page_height"]
        margin_x = context["margin_x"]
        content_width = context["table_width"]
        brand_logo = context["brand_logo"]

        if brand_logo.exists():
            try:
                logo_reader = ImageReader(str(brand_logo))
                source_width, source_height = logo_reader.getSize()
                logo_width = 66.0
                logo_height = logo_width * source_height / source_width
                surface.drawImage(
                    logo_reader,
                    margin_x,
                    page_height - 47.0,
                    width=logo_width,
                    height=logo_height,
                    preserveAspectRatio=False,
                    anchor="sw",
                    mask="auto",
                )
                surface.setFillColor(colors.HexColor("#444444"))
                surface.setFont("Helvetica", 3.15)
                surface.drawCentredString(
                    margin_x + logo_width / 2,
                    page_height - 51.8,
                    "Engenharia, Empreendimentos, Manutenção e Instalações Ltda",
                )
            except Exception:
                pass

        surface.setFillColor(colors.HexColor("#d9d9d9"))
        surface.rect(margin_x, page_height - 72, content_width, 17, fill=1, stroke=0)
        title, title_size = _fit_pdf_text(
            report_title,
            content_width - 12,
            "Helvetica-Bold",
            7.2,
            5.2,
        )
        surface.setFillColor(colors.black)
        surface.setFont("Helvetica-Bold", title_size)
        surface.drawCentredString(page_width / 2, page_height - 66.5, title)
        surface.setFont("Helvetica", 4.8)
        surface.drawCentredString(
            page_width / 2,
            page_height - 82,
            "ENGEMIL - Engenharia, Empreendimentos, Manutenção e Instalações Ltda. "
            "- CNPJ 04.768.702/0001-70",
        )
        surface.drawCentredString(
            page_width / 2,
            page_height - 89,
            "Fórmulas, justificativas e declaração de atendimento",
        )
        surface.setFont("Helvetica", 4.2)
        surface.drawRightString(
            page_width - margin_x,
            page_height - 102,
            f"Página {context['page_number']} de {context['page_count']}",
        )

        normal_style = ParagraphStyle(
            "indices_normal",
            fontName="Helvetica",
            fontSize=8.0,
            leading=10.8,
            alignment=TA_JUSTIFY,
            textColor=colors.black,
            spaceAfter=0,
        )
        center_style = ParagraphStyle(
            "indices_center",
            parent=normal_style,
            alignment=TA_CENTER,
            fontSize=8.1,
            leading=10.5,
        )
        heading_style = ParagraphStyle(
            "indices_heading",
            fontName="Helvetica-Bold",
            fontSize=7.5,
            leading=9.2,
            alignment=TA_LEFT,
            textColor=colors.HexColor("#5A1235"),
            spaceAfter=0,
        )
        card_title_style = ParagraphStyle(
            "indices_card_title",
            fontName="Helvetica-Bold",
            fontSize=7.2,
            leading=8.6,
            alignment=TA_CENTER,
            textColor=colors.HexColor("#5A1235"),
        )
        card_value_style = ParagraphStyle(
            "indices_card_value",
            fontName="Helvetica-Bold",
            fontSize=7.2,
            leading=9.2,
            alignment=TA_CENTER,
            textColor=colors.black,
        )

        def draw_paragraph(text, x, y_top, width, style, gap=5):
            paragraph = Paragraph(
                escape(str(text or "")).replace("\n", "<br/>"),
                style,
            )
            _, height = paragraph.wrap(width, page_height)
            paragraph.drawOn(surface, x, y_top - height)
            return y_top - height - gap

        def draw_heading(text, y_top):
            surface.setFillColor(colors.HexColor("#F2F2F2"))
            heading = Paragraph(escape(str(text)), heading_style)
            _, height = heading.wrap(content_width - 14, 45)
            box_height = max(23.0, height + 9.0)
            surface.roundRect(
                margin_x,
                y_top - box_height,
                content_width,
                box_height,
                3,
                fill=1,
                stroke=0,
            )
            heading.drawOn(
                surface,
                margin_x + 7,
                y_top - box_height + (box_height - height) / 2,
            )
            return y_top - box_height - 7

        total = sum(float(row.get("Valor atual") or 0) for row in rows)
        total_remaining = sum(
            float(row.get("Remanescente total") or 0) for row in rows
        )
        equity = float(parameters.get("equity_value") or 0)
        revenue = float(parameters.get("gross_revenue") or 0)
        index_total = equity * 12 / total if total else 0
        index_remaining = equity * 12 / total_remaining if total_remaining else 0
        variation = ((revenue - total) / revenue * 100) if revenue else 0

        def decimal_br(value):
            return f"{float(value or 0):.2f}".replace(".", ",")

        y = page_height - 118
        y = draw_heading(
            'FÓRMULA PARA ATENDIMENTO AO ITEM "D.1" DA ALÍNEA "D" DO SUBITEM '
            "11.1 DO ITEM 11 DO ANEXO VII-A DA IN SEGES/MP Nº 05/2017",
            y,
        )
        y = draw_paragraph(
            "Declaramos que 1/12 (um doze avos) dos contratos firmados com a "
            "Administração Pública e/ou com a iniciativa privada, vigentes na data "
            "de apresentação da proposta, não é superior ao Patrimônio Líquido da empresa.",
            margin_x,
            y,
            content_width,
            normal_style,
            gap=8,
        )

        card_gap = 8
        card_width = (content_width - card_gap) / 2
        card_height = 58
        card_y = y - card_height
        for index, (card_title, denominator, result) in enumerate((
            ("SOBRE O VALOR TOTAL DOS CONTRATOS", total, index_total),
            ("SOBRE O REMANESCENTE DOS CONTRATOS", total_remaining, index_remaining),
        )):
            card_x = margin_x + index * (card_width + card_gap)
            surface.setFillColor(colors.HexColor("#F7F3F5"))
            surface.setStrokeColor(colors.HexColor("#5A1235"))
            surface.setLineWidth(0.6)
            surface.roundRect(
                card_x,
                card_y,
                card_width,
                card_height,
                5,
                fill=1,
                stroke=1,
            )
            heading = Paragraph(card_title, card_title_style)
            _, heading_height = heading.wrap(card_width - 10, 20)
            heading.drawOn(
                surface,
                card_x + 5,
                card_y + card_height - heading_height - 6,
            )
            status = "ATENDE" if result >= 1 else "NÃO ATENDE"
            formula = Paragraph(
                f"{escape(brl(equity))} x 12 / {escape(brl(denominator))} = "
                f"{decimal_br(result)}<br/><font color='{'#18794E' if result >= 1 else '#B42318'}'>"
                f"{status} ao mínimo de 1,00</font>",
                card_value_style,
            )
            _, formula_height = formula.wrap(card_width - 12, 30)
            formula.drawOn(surface, card_x + 6, card_y + 7, )
        y = card_y - 9

        y = draw_heading(
            'FÓRMULA PARA ATENDIMENTO AO ITEM "D.2" DA ALÍNEA "D" DO SUBITEM '
            "11.1 DO ITEM 11 DO ANEXO VII-A DA IN SEGES/MP Nº 05/2017",
            y,
        )
        y = draw_paragraph(
            "Cálculo demonstrativo da variação percentual do valor total dos "
            "contratos firmados em relação à Receita Bruta discriminada na "
            "Demonstração do Resultado do Exercício (DRE).",
            margin_x,
            y,
            content_width,
            normal_style,
            gap=5,
        )
        y = draw_paragraph(
            f"({brl(revenue)} - {brl(total)}) x 100 / {brl(revenue)} = "
            f"{decimal_br(variation)}%.",
            margin_x,
            y,
            content_width,
            center_style,
            gap=8,
        )
        y = draw_heading("JUSTIFICATIVA DA VARIAÇÃO", y)
        justification = parameters.get("justification_text") or (
            "A divergência observada entre os valores apresentados na Demonstração "
            "do Resultado do Exercício e a relação de contratos firmados decorre da "
            "diferença nos critérios e períodos de reconhecimento das receitas. A DRE "
            "contempla as receitas reconhecidas no exercício, enquanto a relação inclui "
            "instrumentos vigentes com faturamentos distribuídos em exercícios presentes "
            "e futuros."
        )
        draw_paragraph(
            justification,
            margin_x,
            y,
            content_width,
            normal_style,
            gap=0,
        )

        surface.setFillColor(colors.black)
        surface.setFont("Helvetica", 7.4)
        surface.drawCentredString(
            page_width / 2,
            155,
            f"Brasília/DF, {reference_date.strftime('%d/%m/%Y')}.",
        )
        signature_center = page_width / 2
        signature_width = 265
        signature_line_y = 112
        surface.setStrokeColor(colors.black)
        surface.setLineWidth(0.5)
        surface.line(
            signature_center - signature_width / 2,
            signature_line_y,
            signature_center + signature_width / 2,
            signature_line_y,
        )
        signature_lines = [
            (
                "ENGEMIL ENGENHARIA, EMPREENDIMENTOS, MANUTENÇÃO E INSTALAÇÕES LTDA",
                "Helvetica-Bold",
                5.4,
            ),
            (parameters.get("signatory_name") or "", "Helvetica-Bold", 5.3),
            (
                " - ".join(filter(None, [
                    parameters.get("signatory_registration"),
                    (
                        f"CPF {parameters.get('signatory_cpf')}"
                        if parameters.get("signatory_cpf") else ""
                    ),
                ])),
                "Helvetica",
                4.7,
            ),
            (parameters.get("signatory_title") or "", "Helvetica", 4.8),
        ]
        signature_y = signature_line_y - 8
        for text, font_name, font_size in signature_lines:
            if text:
                fitted, fitted_size = _fit_pdf_text(
                    text,
                    signature_width + 50,
                    font_name,
                    font_size,
                    3.8,
                )
                surface.setFont(font_name, fitted_size)
                surface.drawCentredString(signature_center, signature_y, fitted)
                signature_y -= 7

        surface.setStrokeColor(colors.HexColor("#a6a6a6"))
        surface.setLineWidth(0.35)
        surface.line(margin_x, 35.5, page_width - margin_x, 35.5)
        surface.setFillColor(colors.HexColor("#555555"))
        surface.setFont("Helvetica", 3.8)
        surface.drawCentredString(
            page_width / 2,
            29,
            'CRS 503, BLOCO "B", LOJA 05, PARTE SUPERIOR "A" - BRASÍLIA/DF - CEP 70.331-520',
        )
        surface.drawCentredString(
            page_width / 2,
            23.5,
            "(61) 3248-2876 / 3248-5919 / 3248-3410",
        )
        surface.drawCentredString(
            page_width / 2,
            18,
            "gestao.contratos@engemileng.com - https://engemil.com.br/",
        )
        surface.drawRightString(
            page_width - margin_x,
            18,
            reference_date.strftime("%d/%m/%Y"),
        )

    return generate_backlog_pdf(
        rows,
        reference_date=reference_date,
        logo_path=logo_path,
        signatory=None,
        sort_label="Centro de custo - ordem crescente",
        report_title=report_title,
        report_subtitle=(
            "Relação de contratos vigentes - posição em "
            f"{reference_date.strftime('%d/%m/%Y')}"
        ),
        pdf_subject="Declaração de contratos e índices econômico-financeiros",
        force_single_table_page=True,
        additional_page_renderer=draw_details_page,
    )


def generate_contract_dossier(payload, template_path=DEFAULT_TEMPLATE):
    contract = dict(payload.get("contract") or {})
    effective = dict(payload.get("effective") or {})
    doc = _new_document(template_path)
    number = contract.get("contract_number") or "sem número"
    _add_title(
        doc,
        "FICHA CONTRATUAL",
        f"Contrato {number} · {contract.get('client') or 'Contratante não informado'}",
    )

    _add_heading(doc, "Identificação e vigência")
    _add_definition_table(doc, [
        ("Centro de custo", contract.get("cost_center")),
        ("Contratante", contract.get("client")),
        ("Contrato", number),
        ("Edital/licitação", contract.get("bid_number")),
        ("Número do processo", contract.get("process_number")),
        ("UASG", contract.get("uasg")),
        ("Modalidade", contract.get("category")),
        ("Modalidade de licitação", contract.get("procurement_method")),
        (
            "Inscrição no CNO",
            {
                1: "Obrigatória",
                0: "Não aplicável",
                None: "A definir",
            }.get(contract.get("cno_required"), "A definir"),
        ),
        (
            "Regime de faturamento",
            {
                "ONERADO": "Onerado / não desonerado",
                "DESONERADO": "Desonerado",
                "NÃO DEFINIDO": "Não definido",
            }.get(contract.get("tax_regime"), contract.get("tax_regime")),
        ),
        ("Status", contract.get("status")),
        ("Assinatura do contrato", short_date(contract.get("signature_date"))),
        ("Início original", short_date(effective.get("original_start_date") or contract.get("start_date"))),
        ("Fim original", short_date(effective.get("original_end_date") or contract.get("end_date"))),
        ("Valor original", brl(effective.get("original_value") or contract.get("original_value"))),
    ])

    budget_dates = payload.get("budget_dates") or []
    budget_date_summary = "; ".join(
        " — ".join(filter(None, [
            short_date(item.get("reference_date")),
            str(item.get("description") or "").strip(),
        ]))
        for item in budget_dates
        if item.get("reference_date")
    )
    current_and_responsible = [
        ("Instrumento vigente", effective.get("current_instrument")),
        ("Início da vigência atual", short_date(effective.get("current_start_date"))),
        ("Fim da vigência atual", short_date(effective.get("current_end_date"))),
        ("Valor vigente", brl(effective.get("current_value") or contract.get("current_value"))),
        ("Remanescente estimado", brl(effective.get("remaining_value"))),
        ("Situação da vigência", effective.get("lifecycle_status")),
        ("Engenheiro responsável", contract.get("engineer_name")),
        ("E-mail do engenheiro", contract.get("engineer_email")),
        ("Responsável administrativo", contract.get("manager_name")),
        ("E-mail administrativo", contract.get("manager_email")),
        ("Quantidade de empregados", contract.get("employee_count")),
        ("Próxima repactuação", short_date(contract.get("repactuation_date"))),
        ("Data(s) do orçamento", budget_date_summary),
    ]
    if any(_meaningful(value) for _, value in current_and_responsible):
        doc.add_page_break()
        _add_heading(doc, "Vigência atual, responsáveis e controles")
        _add_definition_table(doc, current_and_responsible)
    if _meaningful(contract.get("object")):
        _add_heading(doc, "Objeto do contrato")
        _add_text(doc, contract["object"], justified=True)
    if _meaningful(contract.get("observations")):
        _add_heading(doc, "Observações gerais")
        _add_text(doc, contract["observations"], justified=True)

    amendments = payload.get("amendments") or []
    if amendments:
        doc.add_page_break()
        _add_heading(doc, "Instrumentos contratuais e aditivos")
        rows = []
        for item in amendments:
            instrument = " ".join(
                str(part).strip() for part in (item.get("ordinal"), item.get("kind")) if _meaningful(part)
            )
            rows.append((
                instrument or "Instrumento",
                "\n".join(filter(None, [
                    f"Início: {short_date(item.get('start_date'))}" if item.get("start_date") else "",
                    f"Fim: {short_date(item.get('end_date'))}" if item.get("end_date") else "",
                    f"Duração: {item.get('duration_months')} mês(es)"
                    if item.get("duration_months") is not None else "",
                ])),
                "\n".join(filter(None, [
                    f"Valor: {brl(item.get('value'))}" if _meaningful(item.get("value")) else "",
                    f"Garantia: {item.get('guarantee_status')}" if item.get("guarantee_status") else "",
                    f"ART: {item.get('art_status')}" if item.get("art_status") else "",
                ])),
                "\n".join(
                    str(value) for value in (
                        item.get("description"), item.get("notes"), item.get("justification_text")
                    ) if _meaningful(value)
                ),
            ))
        _add_records_table(
            doc,
            ["Instrumento", "Vigência", "Valor e controles", "Descrição e observações"],
            rows,
            [1.55, 1.85, 2.10, 1.37],
        )

    guarantees = payload.get("guarantees") or []
    if guarantees:
        doc.add_page_break()
        _add_heading(doc, "Garantias contratuais e seguros")
        rows = []
        for item in guarantees:
            rows.append((
                "\n".join(filter(None, [
                    item.get("display_type") or item.get("guarantee_type"),
                    f"Modalidade: {item.get('modality')}" if item.get("modality") else "",
                    f"Apólice/garantia: {item.get('policy_number')}"
                    if item.get("policy_number") else "",
                    f"SUSEP/controle: {item.get('susep_registration')}"
                    if item.get("susep_registration") else "",
                ])),
                "\n".join(filter(None, [
                    f"Instrumento: {item.get('instrument_reference')}"
                    if item.get("instrument_reference") else "",
                    f"Fundamento: {item.get('legal_basis')}" if item.get("legal_basis") else "",
                    f"Situação: {item.get('operational_status')}"
                    if item.get("operational_status") else "",
                ])),
                "\n".join(filter(None, [
                    f"Exigido: {brl(item.get('required_amount'))}",
                    f"Prêmio: {brl(item.get('premium_value'))}"
                    if float(item.get("premium_value") or 0) else "",
                ])),
                "\n".join(filter(None, [
                    f"Emissão: {short_date(item.get('issue_date'))}"
                    if item.get("issue_date") else "",
                    f"Início: {short_date(item.get('start_date'))}"
                    if item.get("start_date") else "",
                    f"Fim: {short_date(item.get('end_date'))}"
                    if item.get("end_date") else "",
                    f"Prazo para apresentação: {short_date(item.get('request_due_date'))}"
                    if item.get("request_due_date") else "",
                ])),
            ))
        _add_records_table(
            doc,
            ["Garantia/seguro", "Referência e situação", "Valores", "Datas e vigência"],
            rows,
            [1.72, 1.85, 1.15, 2.15],
        )
        for item in guarantees:
            coverages = item.get("coverages") or []
            endorsements = item.get("endorsements") or []
            if not coverages and not endorsements and not _meaningful(item.get("object_description")):
                continue
            _add_heading(
                doc,
                f"Detalhamento - {item.get('display_type') or item.get('guarantee_type') or 'garantia'}",
            )
            if _meaningful(item.get("object_description")):
                _add_text(
                    doc,
                    f"Objeto e obrigações cobertas: {item['object_description']}",
                    justified=True,
                )
            if coverages:
                coverage_rows = [(
                    coverage.get("coverage_name"),
                    brl(coverage.get("insured_limit")),
                    "\n".join(filter(None, [
                        f"Início: {short_date(coverage.get('start_date'))}"
                        if coverage.get("start_date") else "",
                        f"Fim: {short_date(coverage.get('end_date'))}"
                        if coverage.get("end_date") else "",
                    ])),
                    "\n".join(filter(None, [
                        f"Franquia/POS: {coverage.get('deductible')}"
                        if coverage.get("deductible") else "",
                        str(coverage.get("notes") or ""),
                    ])),
                ) for coverage in coverages]
                _add_records_table(
                    doc,
                    ["Cobertura", "LMI", "Vigência", "Franquia/POS e observações"],
                    coverage_rows,
                    [1.62, 1.15, 1.70, 2.40],
                )
            if endorsements:
                endorsement_rows = [(
                    "\n".join(filter(None, [
                        endorsement.get("movement_type"),
                        endorsement.get("endorsement_number"),
                    ])),
                    "\n".join(filter(None, [
                        f"Emissão: {short_date(endorsement.get('issue_date'))}"
                        if endorsement.get("issue_date") else "",
                        f"Novo fim: {short_date(endorsement.get('new_end_date'))}"
                        if endorsement.get("new_end_date") else "",
                    ])),
                    "\n".join(filter(None, [
                        f"Ajuste de prêmio: {brl(endorsement.get('premium_adjustment'))}"
                        if float(endorsement.get("premium_adjustment") or 0) else "",
                    ])),
                    "\n".join(filter(None, [
                        f"Situação: {endorsement.get('request_status')}"
                        if endorsement.get("request_status") else "",
                        str(endorsement.get("description") or ""),
                        str(endorsement.get("notes") or ""),
                    ])),
                ) for endorsement in endorsements]
                _add_records_table(
                    doc,
                    ["Movimentação", "Emissão e vigência", "Prêmio", "Situação e descrição"],
                    endorsement_rows,
                    [1.40, 1.75, 1.30, 2.42],
                )

    bdis = payload.get("bdis") or []
    if bdis:
        doc.add_page_break()
        _add_heading(doc, "Custos Indiretos, Tributos e Lucro — BDI")
        for item_index, item in enumerate(bdis):
            if item_index:
                doc.add_page_break()
                _add_heading(
                    doc,
                    f"Custos Indiretos, Tributos e Lucro — {item.get('name') or 'BDI'}",
                )
            method = item.get("calculation_method")
            method_label = (
                "Soma direta" if method == "SOMA_DIRETA" else "Fórmula composta"
            )
            regime = item.get("effective_tax_regime") or item.get("tax_regime")
            regime_label = {
                "ONERADO": "Onerado / não desonerado",
                "DESONERADO": "Desonerado",
                "NÃO DEFINIDO": "Não definido",
            }.get(regime, regime)
            _add_definition_table(doc, [
                ("BDI", item.get("name")),
                ("Referência/aplicação", item.get("reference_name")),
                ("Regime", regime_label),
                ("Método", method_label),
                ("BDI calculado", percent(item.get("calculated_percentage"))),
            ])
            if method == "SOMA_DIRETA":
                rows = [
                    ("Custos indiretos", percent(item.get("indirect_costs"), 4)),
                    ("Lucro", percent(item.get("profit"), 4)),
                ]
            else:
                rows = [
                    ("Administração central — AC", percent(item.get("central_administration"), 4)),
                    ("Seguros — S", percent(item.get("insurance"), 4)),
                    ("Riscos — R", percent(item.get("risks"), 4)),
                    ("Garantias — G", percent(item.get("guarantees"), 4)),
                    ("Outros custos indiretos", percent(item.get("other_indirect_costs"), 4)),
                    (
                        "Subtotal AC + S + R + G + outros",
                        percent(item.get("composed_indirect_total"), 4),
                    ),
                    ("Despesas financeiras — DF", percent(item.get("financial_expenses"), 4)),
                    ("Lucro — L", percent(item.get("profit"), 4)),
                ]
            rows.extend([
                ("PIS", percent(item.get("pis"), 4)),
                ("COFINS", percent(item.get("cofins"), 4)),
                ("ISS", percent(item.get("iss"), 4)),
                ("CPRB", percent(item.get("cprb"), 4)),
                ("Outros tributos", percent(item.get("other_taxes"), 4)),
                ("Total dos tributos — T", percent(item.get("tax_total"), 4)),
                ("BDI calculado", percent(item.get("calculated_percentage"))),
            ])
            _add_records_table(
                doc,
                ["Parcela", "Percentual"],
                rows,
                [5.30, 1.57],
            )
            if _meaningful(item.get("notes")):
                _add_text(doc, f"Observações: {item['notes']}", justified=True)

    ata_contracts = payload.get("ata_contracts") or []
    if ata_contracts:
        _add_heading(doc, "Contratos decorrentes da ATA")
        for ata in ata_contracts:
            _add_definition_table(doc, [
                ("Contrato", ata.get("contract_number")),
                ("Processo", ata.get("process_number")),
                ("Contratante", ata.get("client")),
                ("Objeto", ata.get("object")),
                ("Assinatura", short_date(ata.get("signature_date"))),
                ("Início", short_date(ata.get("start_date"))),
                ("Fim", short_date(ata.get("end_date"))),
                ("Valor original", brl(ata.get("original_value"))),
                ("Valor atual", brl(ata.get("current_value"))),
                ("Status", ata.get("status")),
                ("Responsável", ata.get("responsible_name")),
                ("E-mail", ata.get("responsible_email")),
                ("Observações", ata.get("notes")),
            ])
            ata_amendments = ata.get("amendments") or []
            if ata_amendments:
                rows = []
                for item in ata_amendments:
                    rows.append((
                        " ".join(filter(None, [str(item.get("ordinal") or ""), str(item.get("kind") or "")])).strip(),
                        "\n".join(filter(None, [
                            f"Início: {short_date(item.get('start_date'))}" if item.get("start_date") else "",
                            f"Fim: {short_date(item.get('end_date'))}" if item.get("end_date") else "",
                        ])),
                        brl(item.get("value")) if _meaningful(item.get("value")) else "",
                        "\n".join(
                            str(value) for value in (item.get("description"), item.get("notes"))
                            if _meaningful(value)
                        ),
                    ))
                _add_records_table(
                    doc,
                    ["Aditivo", "Vigência", "Valor", "Descrição e observações"],
                    rows,
                    [1.25, 1.45, 1.30, 2.87],
                )

    unions = payload.get("unions") or []
    if unions:
        _add_heading(doc, "Sindicatos e datas-base")
        rows = []
        for item in unions:
            rows.append((
                "\n".join(filter(None, [
                    f"Sindicato: {item.get('union_name')}" if item.get("union_name") else "",
                    f"CCT: {item.get('collective_agreement')}" if item.get("collective_agreement") else "",
                    f"Categoria: {item.get('category_name')}" if item.get("category_name") else "",
                ])),
                "\n".join(filter(None, [
                    f"Data-base: {short_date(item.get('base_date'))}" if item.get("base_date") else "",
                    f"Mês-base: {item.get('base_month')}" if item.get("base_month") else "",
                    f"Próxima repactuação: {short_date(item.get('next_repactuation'))}"
                    if item.get("next_repactuation") else "",
                ])),
                item.get("instrument_reference") or "",
                item.get("notes") or "",
            ))
        _add_records_table(
            doc,
            ["Sindicato/CCT", "Data-base e repactuação", "Instrumento de referência", "Observações"],
            rows,
            [1.85, 1.55, 1.55, 1.92],
        )

    positions = payload.get("positions") or []
    if positions:
        _add_heading(doc, "Equipe, cargos e benefícios")
        rows = []
        for item in positions:
            benefit_lines = []
            for benefit in item.get("benefits") or []:
                description = f" - {benefit.get('description')}" if benefit.get("description") else ""
                benefit_lines.append(
                    f"{benefit.get('benefit_type') or 'Benefício'}: "
                    f"{brl(benefit.get('monthly_value'))}{description}"
                )
            rows.append((
                "\n".join(filter(None, [
                    f"Cargo: {item.get('title')}" if item.get("title") else "",
                    f"Quantidade: {item.get('quantity')}" if _meaningful(item.get("quantity")) else "",
                ])),
                "\n".join(filter(None, [
                    f"Salário-base: {brl(item.get('base_salary'))}",
                    f"Periculosidade: {item.get('hazard_percent')}% / {brl(item.get('additional_value'))}"
                    if float(item.get("hazard_percent") or 0) else "",
                    f"Insalubridade: {item.get('unhealthy_percent')}% / ano-base "
                    f"{item.get('unhealthy_base_year')}"
                    if float(item.get("unhealthy_percent") or 0) else "",
                    f"Benefícios mensais: {brl(item.get('monthly_benefits'))}"
                    if float(item.get("monthly_benefits") or 0) else "",
                ])),
                "\n".join(benefit_lines),
                "\n".join(filter(None, [
                    f"Sindicato/CCT: {item.get('union_reference')}"
                    if item.get("union_reference") else "",
                    str(item.get("notes") or ""),
                ])),
            ))
        _add_records_table(
            doc,
            ["Cargo e equipe", "Remuneração e adicionais", "Benefícios detalhados", "Referência e observações"],
            rows,
            [1.35, 1.70, 1.85, 1.97],
        )

    obligations = payload.get("obligations") or []
    if obligations:
        _add_heading(doc, "Prazos e obrigações")
        rows = []
        for item in obligations:
            rows.append((
                "\n".join(filter(None, [
                    f"Obrigação: {item.get('title')}" if item.get("title") else "",
                    f"Categoria: {item.get('category')}" if item.get("category") else "",
                    f"Status: {item.get('status')}" if item.get("status") else "",
                ])),
                "\n".join(filter(None, [
                    f"Vencimento: {short_date(item.get('due_date'))}" if item.get("due_date") else "",
                    f"Prioridade: {item.get('priority')}" if item.get("priority") else "",
                    f"Recorrência: {item.get('recurrence')}" if item.get("recurrence") else "",
                ])),
                "\n".join(filter(None, [
                    f"Responsável: {item.get('responsible_name')}" if item.get("responsible_name") else "",
                    f"E-mail: {item.get('responsible_email')}" if item.get("responsible_email") else "",
                    f"Cópia/grupo: {item.get('copy_emails')}" if item.get("copy_emails") else "",
                ])),
                "\n".join(filter(None, [
                    f"Antecedência: {item.get('advance_days')} dias"
                    if _meaningful(item.get("advance_days")) else "",
                    f"Lembretes: a cada {item.get('reminder_frequency_days')} dias"
                    if item.get("notification_enabled") else "Notificações desabilitadas",
                    str(item.get("notes") or ""),
                ])),
            ))
        _add_records_table(
            doc,
            ["Obrigação", "Prazo", "Responsável e destinatários", "Alertas e observações"],
            rows,
            [1.50, 1.30, 2.30, 1.77],
        )

    arts = payload.get("arts") or []
    if arts:
        _add_heading(doc, "Anotações de Responsabilidade Técnica - ART")
        rows = []
        for item in arts:
            rows.append((
                "\n".join(filter(None, [
                    f"Profissional: {item.get('professional_display_name') or item.get('professional_name')}"
                    if item.get("professional_display_name") or item.get("professional_name") else "",
                    f"Título: {item.get('professional_title')}"
                    if item.get("professional_title") else "",
                    f"Registro: {item.get('professional_registration')}"
                    if item.get("professional_registration") else "",
                ])),
                "\n".join(filter(None, [
                    item.get("art_number") or "",
                    (
                        f"Instrumento: {item.get('instrument_reference')}"
                        if item.get("instrument_reference") else ""
                    ),
                ])),
                "\n".join(filter(None, [
                    f"Emissão: {short_date(item.get('issue_date'))}" if item.get("issue_date") else "",
                    f"Término: {short_date(item.get('end_date'))}" if item.get("end_date") else "",
                    f"Status: {item.get('status')}" if item.get("status") else "",
                ])),
                "\n".join(
                    str(value) for value in (item.get("description"), item.get("notes"))
                    if _meaningful(value)
                ),
            ))
        _add_records_table(
            doc,
            ["Profissional", "ART e instrumento", "Vigência e status", "Descrição e observações"],
            rows,
            [1.75, 1.35, 1.45, 2.32],
        )

    cnos = payload.get("cnos") or []
    if cnos:
        _add_heading(doc, "Cadastro Nacional de Obras — CNO")
        rows = []
        for item in cnos:
            rows.append((
                "\n".join(filter(None, [
                    item.get("registration_number") or "",
                    (
                        f"Vinculado a: {item.get('ata_contract_reference')}"
                        if item.get("ata_contract_reference") else ""
                    ),
                ])),
                "\n".join(filter(None, [
                    f"Cadastramento: {short_date(item.get('registration_date'))}"
                    if item.get("registration_date") else "",
                    f"Início da responsabilidade: "
                    f"{short_date(item.get('responsibility_start_date'))}"
                    if item.get("responsibility_start_date") else "",
                ])),
                item.get("work_area") or "",
                item.get("notes") or "",
            ))
        _add_records_table(
            doc,
            [
                "Número de inscrição", "Datas do cadastro e responsabilidade",
                "Área de atuação da obra", "Observações",
            ],
            rows,
            [1.30, 1.85, 2.15, 1.57],
        )

    documents = payload.get("documents") or []
    if documents:
        _add_heading(doc, "Documentos anexados")
        rows = [
            (
                "\n".join(filter(None, [
                    f"Categoria: {item.get('category')}" if item.get("category") else "",
                    f"Vinculado a: {item.get('association')}" if item.get("association") else "",
                ])),
                "\n".join(filter(None, [
                    f"Título: {item.get('title')}" if item.get("title") else "",
                    f"Arquivo: {item.get('filename')}" if item.get("filename") else "",
                ])),
                short_datetime(item.get("uploaded_at")),
            )
            for item in documents
        ]
        _add_records_table(
            doc,
            ["Categoria e vínculo", "Documento", "Inclusão no sistema"],
            rows,
            [1.75, 3.37, 1.75],
        )

    generated = payload.get("generated_documents") or []
    if generated:
        _add_heading(doc, "Documentos padronizados gerados")
        rows = [
            (
                "\n".join(filter(None, [
                    f"Número: {item.get('document_number')}" if item.get("document_number") else "",
                    f"Modelo: {item.get('template_name')}" if item.get("template_name") else "",
                    f"Status: {item.get('status')}" if item.get("status") else "",
                ])),
                "\n".join(filter(None, [
                    f"Destinatário: {item.get('recipient')}" if item.get("recipient") else "",
                    f"Assunto: {item.get('subject')}" if item.get("subject") else "",
                ])),
                "\n".join(filter(None, [
                    f"Gerado por: {item.get('created_by_name')}" if item.get("created_by_name") else "",
                    f"Gerado em: {short_datetime(item.get('created_at'))}"
                    if item.get("created_at") else "",
                    f"Encaminhado em: {short_datetime(item.get('sent_at'))}"
                    if item.get("sent_at") else "",
                ])),
                item.get("notes") or "",
            )
            for item in generated
        ]
        _add_records_table(
            doc,
            ["Documento", "Destinatário e assunto", "Rastreabilidade", "Observações"],
            rows,
            [1.60, 1.85, 1.80, 1.62],
        )

    return _save_bytes(doc, "FICHA CONTRATUAL")
