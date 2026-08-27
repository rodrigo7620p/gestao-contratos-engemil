import os
import tempfile
from datetime import date, timedelta
from decimal import Decimal
from io import BytesIO
from pathlib import Path

tmp = tempfile.TemporaryDirectory()
os.environ["GESTAO_DB_PATH"] = str(Path(tmp.name) / "test.db")
os.environ["GESTAO_UPLOAD_DIR"] = str(Path(tmp.name) / "uploads")

from db import (
    archive_expired_contracts,
    authenticate,
    execute,
    init_db,
    next_document_sequence,
    query,
    refresh_contract_lifecycle,
)
from alerts import obligation_notification_text, process_repactuation_alerts
from bdi import calculate_bdi, composed_indirect_total, tax_total
from declaration import generate_declaration
from document_factory import (
    extract_placeholders,
    format_document_number,
    generate_document,
    resolve_project_path,
)
from importer import import_workbook
from notifications import normalize_recipients
from portfolio import annual_allocation, backlog_rows, remaining_value, workbook_bytes
from reports import generate_backlog_pdf, generate_contract_dossier
from totp import code, new_secret, provisioning_uri, verify
import pandas as pd
from docx import Document


def run():
    init_db()
    assert authenticate("admin@engemil.local", "Alterar@123")
    filename = "02_ANALISE_CRITICA_DE_CONTRATOS(2).xlsx"
    base = Path(__file__).resolve().parent
    source = next(
        (
            candidate for candidate in (
                base.parent / "upload" / filename,
                base.parent.parent / "upload" / filename,
            )
            if candidate.is_file()
        ),
        None,
    )
    if source:
        result = import_workbook(source, replace_amendments=True)
        assert result["contracts"] > 50, result
        assert query("SELECT COUNT(*) n FROM amendments")[0]["n"] > 50
        assert not query(
            """SELECT id FROM amendments
            WHERE UPPER(TRIM(COALESCE(ordinal,'')))
                IN ('INICIAL','CONTRATO INICIAL')
            AND UPPER(TRIM(COALESCE(kind,'')))
                IN ('CONTRATO','CONTRATO INICIAL')"""
        )
        dataprev = query(
            "SELECT * FROM contracts WHERE cost_center=?", ("01.04.00288",)
        )[0]
        assert dataprev["client"].startswith("EMPRESA DE TECNOLOGIA")
        assert round(dataprev["current_value"], 2) == 86353294.29
    else:
        dataprev_id = execute(
            """INSERT INTO contracts(
            cost_center,client,contract_number,category,object,start_date,end_date,
            original_value,current_value,status)
            VALUES(?,?,?,?,?,?,?,?,?,?)""",
            (
                "01.04.00288",
                "EMPRESA DE TECNOLOGIA E INFORMAÇÕES DA PREVIDÊNCIA S.A. - DATAPREV",
                "01.038342.2024",
                "REFORMA",
                "Reforma predial de teste, cadastrada pelo caminho de contingência "
                "do test_core.py (sem a planilha legada de origem).",
                "2026-01-01",
                "2027-12-31",
                86353294.29,
                86353294.29,
                "ATIVO",
            ),
        )
        dataprev = query("SELECT * FROM contracts WHERE id=?", (dataprev_id,))[0]
        result = None
    secret = new_secret()
    assert verify(secret, code(secret))
    assert not verify(secret, "000000") or code(secret) == "000000"
    assert provisioning_uri(secret, "teste@engemil.com").startswith("otpauth://totp/")
    user_columns = {r["name"] for r in query("PRAGMA table_info(users)")}
    contract_columns = {r["name"] for r in query("PRAGMA table_info(contracts)")}
    document_columns = {r["name"] for r in query("PRAGMA table_info(documents)")}
    obligation_columns = {r["name"] for r in query("PRAGMA table_info(obligations)")}
    assert {"totp_secret", "totp_enabled", "preferred_theme"} <= user_columns
    assert {"archived", "archived_at", "archived_by", "tax_regime"} <= contract_columns
    assert {"ata_contract_id", "ata_amendment_id"} <= document_columns
    assert {
        "copy_emails", "notification_enabled", "reminder_frequency_days",
        "last_reminder_at", "reminder_count",
    } <= obligation_columns
    assert normalize_recipients(
        "gestor@engemil.com; engenharia@engemil.com, gestor@engemil.com"
    ) == ["gestor@engemil.com", "engenharia@engemil.com"]
    assert query("SELECT name FROM sqlite_master WHERE type='table' AND name='ata_contracts'")
    assert query("SELECT name FROM sqlite_master WHERE type='table' AND name='contract_bdis'")
    assert query("SELECT COUNT(*) n FROM company_document_templates")[0]["n"] == 3
    assert query("SELECT COUNT(*) n FROM company_signatories")[0]["n"] == 2
    assert next_document_sequence("OFICIO", 2026) == 1
    assert next_document_sequence("OFICIO", 2026) == 2
    assert format_document_number("OFICIO", 2, "10/2026", "AGU", 2026).startswith(
        "OF-2026-0002/ENGEMIL/DCONT"
    )
    old_contract_id = execute(
        """INSERT INTO contracts(cost_center,client,contract_number,start_date,end_date,status)
        VALUES(?,?,?,?,?,'ATIVO')""",
        (
            "TESTE.ARQUIVO", "Cliente Teste", "CT-ANTIGO",
            (date.today() - timedelta(days=400)).isoformat(),
            (date.today() - timedelta(days=31)).isoformat(),
        ),
    )
    assert old_contract_id in archive_expired_contracts()
    assert query("SELECT archived FROM contracts WHERE id=?", (old_contract_id,))[0]["archived"] == 1
    execute(
        """INSERT INTO amendments(contract_id,ordinal,kind,start_date,end_date,value)
        VALUES(?,?,'TERMO ADITIVO',?,?,?)""",
        (
            old_contract_id, "1º", date.today().isoformat(),
            (date.today() + timedelta(days=365)).isoformat(), 1000,
        ),
    )
    assert refresh_contract_lifecycle(old_contract_id) == "ATIVO"
    assert query("SELECT archived FROM contracts WHERE id=?", (old_contract_id,))[0]["archived"] == 0
    ata_contract_id = execute(
        """INSERT INTO ata_contracts(
        ata_id,contract_number,client,start_date,end_date,original_value,current_value)
        VALUES(?,?,?,?,?,?,?)""",
        (
            dataprev["id"], "ATA-CT-01", "Cliente ATA", date.today().isoformat(),
            (date.today() + timedelta(days=365)).isoformat(), 5000, 5000,
        ),
    )
    ata_amendment_id = execute(
        """INSERT INTO ata_contract_amendments(
        ata_contract_id,ordinal,kind,end_date,value) VALUES(?,?,'TERMO ADITIVO',?,?)""",
        (ata_contract_id, "1º", (date.today() + timedelta(days=730)).isoformat(), 6500),
    )
    assert ata_amendment_id
    union_id = execute(
        """INSERT INTO contract_unions(contract_id,union_name,collective_agreement,base_month,next_repactuation)
        VALUES(?,?,?,?,date('now','+60 days'))""",
        (dataprev["id"], "Sindicato Teste", "CCT 2026", 5),
    )
    position_id = execute(
        """INSERT INTO contract_positions(contract_id,title,quantity,base_salary,monthly_benefits,union_id)
        VALUES(?,?,?,?,?,?)""",
        (dataprev["id"], "Eletricista", 2, 3000, 800, union_id),
    )
    execute(
        "INSERT INTO position_benefits(position_id,benefit_type,description,monthly_value) VALUES(?,?,?,?)",
        (position_id, "PLANO DE SAÚDE", "Teste", 250),
    )
    art_id = execute(
        "INSERT INTO arts(contract_id,professional_name,art_number,status) VALUES(?,?,?,'ATIVA')",
        (dataprev["id"], "Profissional Teste", "ART-001"),
    )
    execute(
        """INSERT INTO documents(contract_id,union_id,art_id,category,title,filename,stored_path)
        VALUES(?,?,?,?,?,?,?)""",
        (dataprev["id"], union_id, art_id, "TESTE", "Documento", "teste.pdf", "/tmp/teste.pdf"),
    )
    assert query("SELECT SUM(quantity) total FROM contract_positions")[0]["total"] == 2
    assert query("SELECT SUM(monthly_value) total FROM position_benefits")[0]["total"] == 250
    assert query("SELECT COUNT(*) n FROM arts")[0]["n"] == 1
    alert_result = process_repactuation_alerts()
    assert alert_result["created"] == 1
    assert query("SELECT COUNT(*) n FROM obligations WHERE category='REPACTUAÇÃO'")[0]["n"] == 1
    subject, body = obligation_notification_text({
        "due_date": (date.today() - timedelta(days=2)).isoformat(),
        "title": "Providenciar garantia",
        "client": "Órgão Teste",
        "contract_number": "10/2026",
        "cost_center": "01.01.00001",
        "category": "GARANTIA",
        "priority": "ALTA",
        "responsible_name": "Responsável",
        "notes": "Encaminhar apólice.",
    })
    assert "VENCIDA HÁ 2 DIA(S)" in subject
    assert "responda a este e-mail com a confirmação e as evidências" in body
    assert annual_allocation("2026-01-01", "2026-12-31", 1200, 2026) == 1200
    assert remaining_value("2026-01-01", "2026-12-31", 1200)
    direct_bdi = {
        "calculation_method": "SOMA_DIRETA",
        "indirect_costs": 3.20,
        "profit": 4.00,
        "pis": 0.65,
        "cofins": 3.00,
        "iss": 5.00,
        "cprb": 4.50,
        "other_taxes": 0,
    }
    assert calculate_bdi(direct_bdi) == Decimal("20.3500")
    desonerated_bdi = {
        "calculation_method": "FORMULA_COMPOSTA",
        "rounding_method": "TRUNCAR_4",
        "central_administration": 5.50,
        "insurance": 1.00,
        "risks": 1.27,
        "guarantees": 0,
        "other_indirect_costs": 0,
        "financial_expenses": 1.39,
        "profit": 8.95,
        "pis": 0.65,
        "cofins": 3.00,
        "iss": 2.00,
        "cprb": 4.50,
        "other_taxes": 0,
    }
    assert composed_indirect_total(desonerated_bdi) == Decimal("7.77")
    assert tax_total(desonerated_bdi) == Decimal("10.15")
    assert calculate_bdi(desonerated_bdi) == Decimal("32.49")
    onerated_bdi = {
        **desonerated_bdi,
        "rounding_method": "ARREDONDAR_2",
        "central_administration": 4.00,
        "insurance": 0.80,
        "financial_expenses": 1.23,
        "profit": 7.40,
        "cprb": 0,
    }
    assert calculate_bdi(onerated_bdi) == Decimal("22.23")
    bdi_id = execute(
        """INSERT INTO contract_bdis(
        contract_id,name,reference_name,tax_regime,calculation_method,rounding_method,
        central_administration,insurance,risks,guarantees,financial_expenses,profit,
        pis,cofins,iss,cprb)
        VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",
        (
            dataprev["id"], "BDI 1", "Mão de obra", "DESONERADO",
            "FORMULA_COMPOSTA", "TRUNCAR_4", 5.50, 1.00, 1.27, 0,
            1.39, 8.95, 0.65, 3.00, 2.00, 4.50,
        ),
    )
    assert query(
        "SELECT id FROM contract_bdis WHERE id=? AND contract_id=?",
        (bdi_id, dataprev["id"]),
    )
    backlog = backlog_rows([dict(dataprev)], 2026, 6)
    assert len(backlog) == 1 and "2031" in backlog[0]
    exported = workbook_bytes({"Contratos": pd.DataFrame(backlog)})
    assert exported[:2] == b"PK"
    declaration_contracts = []
    for contract in query("SELECT * FROM contracts WHERE archived=0 ORDER BY client"):
        item = dict(contract)
        item["current_instrument"] = "Contrato"
        item["remaining_value"] = remaining_value(
            item["start_date"], item["end_date"], item["current_value"]
        )
        declaration_contracts.append(item)
    parameters = dict(query("SELECT * FROM financial_parameters WHERE id=1")[0])
    declaration_bytes = generate_declaration(declaration_contracts, parameters)
    assert declaration_bytes[:2] == b"PK"
    declaration_doc = Document(BytesIO(declaration_bytes))
    assert declaration_doc.styles["Normal"].font.name == "Calibri"
    assert round(declaration_doc.styles["Normal"].font.size.pt) == 11
    assert all(
        section.page_height > section.page_width
        for section in declaration_doc.sections
    )
    assert "ENGEMIL ENGENHARIA" in "\n".join(
        paragraph.text for paragraph in declaration_doc.paragraphs
    )
    template = dict(query(
        "SELECT * FROM company_document_templates WHERE document_type='OFICIO'"
    )[0])
    template_path = resolve_project_path(template["generation_path"])
    placeholders = extract_placeholders(template_path)
    assert "{{NUMERO_OFICIO}}" in placeholders
    generated_path = Path(tmp.name) / "oficio_teste.docx"
    generate_document(
        template_path,
        generated_path,
        {token: f"TESTE {token[2:-2]}" for token in placeholders},
    )
    assert generated_path.exists()
    assert not extract_placeholders(generated_path)
    generated_doc = Document(generated_path)
    assert all(section.page_height > section.page_width for section in generated_doc.sections)
    generated_id = execute(
        """INSERT INTO generated_company_documents(
        template_id,contract_id,document_number,recipient,subject,status,
        docx_filename,docx_path,fields_json,created_by,notes)
        VALUES(?,?,?,?,?,'GERADO',?,?,?,?,?)""",
        (
            template["id"], dataprev["id"], "OF-TESTE", dataprev["client"],
            "Documento de teste", generated_path.name, str(generated_path), "{}",
            authenticate("admin@engemil.local", "Alterar@123")["id"], "Somente Word",
        ),
    )
    generated_record = dict(query(
        "SELECT * FROM generated_company_documents WHERE id=?", (generated_id,)
    )[0])
    assert generated_record["docx_filename"] == generated_path.name
    assert generated_record["pdf_filename"] is None
    assert generated_record["pdf_path"] is None
    backlog_pdf = generate_backlog_pdf(backlog)
    assert backlog_pdf.startswith(b"%PDF-")
    assert b"Backlog ENGEMIL" in backlog_pdf
    dossier = generate_contract_dossier({
        "contract": dict(dataprev),
        "effective": {
            "original_start_date": dataprev["start_date"],
            "original_end_date": dataprev["end_date"],
            "current_start_date": dataprev["start_date"],
            "current_end_date": dataprev["end_date"],
            "current_value": dataprev["current_value"],
            "original_value": dataprev["original_value"],
            "current_instrument": "Contrato",
            "remaining_value": remaining_value(
                dataprev["start_date"], dataprev["end_date"], dataprev["current_value"]
            ),
            "lifecycle_status": "VIGENTE",
        },
        "bdis": [
            {
                **dict(query(
                    "SELECT * FROM contract_bdis WHERE id=?", (bdi_id,)
                )[0]),
                "tax_total": float(tax_total(desonerated_bdi)),
                "composed_indirect_total": float(
                    composed_indirect_total(desonerated_bdi)
                ),
                "calculated_percentage": float(calculate_bdi(desonerated_bdi)),
                "effective_tax_regime": "DESONERADO",
            },
            {
                **direct_bdi,
                "name": "BDI 2",
                "reference_name": "Materiais",
                "tax_regime": "DESONERADO",
                "effective_tax_regime": "DESONERADO",
                "tax_total": float(tax_total(direct_bdi)),
                "composed_indirect_total": 0,
                "calculated_percentage": float(calculate_bdi(direct_bdi)),
                "notes": "Composição simplificada por soma direta.",
            },
        ],
        "positions": [{
            **dict(query("SELECT * FROM contract_positions WHERE id=?", (position_id,))[0]),
            "benefits": [
                dict(row) for row in query(
                    "SELECT * FROM position_benefits WHERE position_id=?", (position_id,)
                )
            ],
        }],
        "arts": [dict(query("SELECT * FROM arts WHERE id=?", (art_id,))[0])],
    })
    dossier_doc = Document(BytesIO(dossier))
    assert all(section.page_height > section.page_width for section in dossier_doc.sections)
    dossier_text = "\n".join(paragraph.text for paragraph in dossier_doc.paragraphs)
    assert "FICHA CONTRATUAL" in dossier_text
    assert "Objeto do contrato" in dossier_text
    assert "Custos Indiretos, Tributos e Lucro" in dossier_text
    dossier_qa_output = Path(tempfile.gettempdir()) / "ficha_contratual_bdi_qa.docx"
    dossier_qa_output.write_bytes(dossier)
    qa_output = Path(tempfile.gettempdir()) / "declaracao_contratos_qa.docx"
    qa_output.write_bytes(declaration_bytes)
    print(result)
    print("Testes concluídos com sucesso.")


if __name__ == "__main__":
    run()
