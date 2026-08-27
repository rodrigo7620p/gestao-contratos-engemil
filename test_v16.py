import os
import tempfile
from datetime import date
from io import BytesIO
from pathlib import Path

from docx import Document

temporary = tempfile.TemporaryDirectory()
os.environ["GESTAO_DB_PATH"] = str(Path(temporary.name) / "test_v16.db")
os.environ["GESTAO_UPLOAD_DIR"] = str(Path(temporary.name) / "uploads")

from contract_utils import (  # noqa: E402
    agency_document_fields,
    contract_duration_months,
    extract_agency_acronym,
    normalize_agency_name,
    parse_brazilian_number,
)
from db import execute, init_db, query  # noqa: E402
from document_factory import (  # noqa: E402
    format_document_number,
    generate_document,
)
from reports import generate_contract_dossier  # noqa: E402


def run():
    init_db()
    assert parse_brazilian_number("R$ 18.740.995,83") == 18740995.83
    assert parse_brazilian_number("4163.00") == 4163
    client = "FUNDO NACIONAL DE DESENVOLVIMENTO DA EDUCAÇÃO - FNDE"
    assert extract_agency_acronym(client) == "FNDE"
    assert agency_document_fields(client) == (
        "FUNDO NACIONAL DE DESENVOLVIMENTO DA EDUCAÇÃO",
        "FNDE",
    )
    assert normalize_agency_name(f"{client} - FNDE") == client
    assert contract_duration_months("2026-07-27", "2027-04-01") == 8
    assert contract_duration_months("2026-01-01", "2027-01-01") == 12
    assert format_document_number(
        "OFICIO", 1, "337/2026", "FNDE", 2026
    ).endswith("/337/2026/FNDE")
    agency_name, agency_acronym = agency_document_fields(client)
    document_path = Path(temporary.name) / "oficio_sigla.docx"
    generate_document(
        Path("templates/company_documents/MODELO_OFICIO_ENGEMIL-2026.docx"),
        document_path,
        {
            "ORGAO": agency_name,
            "SIGLA": agency_acronym,
            "CONTRATO": "337/2026",
            "NUMERO_OFICIO": "OF-2026-0001/ENGEMIL/DCONT/337/2026/FNDE",
        },
    )
    generated = Document(document_path)
    generated_text = "\n".join(paragraph.text for paragraph in generated.paragraphs)
    assert "FUNDO NACIONAL DE DESENVOLVIMENTO DA EDUCAÇÃO - FNDE" in generated_text
    assert "FNDE - FNDE" not in generated_text

    contract_id = execute(
        """INSERT INTO contracts(
        cost_center,client,contract_number,start_date,end_date,original_value,
        current_value,status) VALUES(?,?,?,?,?,?,?,'ATIVO')""",
        (
            "TESTE.V16", client, "337/2026", "2026-08-01", "2027-08-01",
            100000, 100000,
        ),
    )
    position_id = execute(
        """INSERT INTO contract_positions(
        contract_id,title,quantity,base_salary) VALUES(?,?,?,?)""",
        (contract_id, "ENGENHEIRO CIVIL", 1, parse_brazilian_number("R$ 12.500,00")),
    )
    assert query(
        "SELECT base_salary FROM contract_positions WHERE id=?", (position_id,)
    )[0]["base_salary"] == 12500
    amendment_id = execute(
        """INSERT INTO amendments(
        contract_id,ordinal,kind,start_date,end_date,duration_months)
        VALUES(?,?,?,?,?,?)""",
        (
            contract_id, "1º", "TERMO ADITIVO", "2027-08-01", "2028-08-01",
            contract_duration_months("2027-08-01", "2028-08-01"),
        ),
    )
    assert query(
        "SELECT duration_months FROM amendments WHERE id=?", (amendment_id,)
    )[0]["duration_months"] == 12
    art_id = execute(
        """INSERT INTO arts(
        contract_id,professional_name,professional_title,
        professional_registration,art_number,status)
        VALUES(?,?,?,?,?,'ATIVA')""",
        (
            contract_id, "Profissional Teste", "Engenheiro Civil",
            "CREA-DF 00000", "ART-TESTE",
        ),
    )
    cno_id = execute(
        """INSERT INTO contract_cnos(
        contract_id,registration_number,registration_date,
        responsibility_start_date,work_area)
        VALUES(?,?,?,?,?)""",
        (
            contract_id, "90.000.00000/00", "2026-08-01", "2026-08-05",
            "Manutenção predial",
        ),
    )
    assert art_id and cno_id
    assert "professional_title" in {
        row["name"] for row in query("PRAGMA table_info(arts)")
    }
    assert "cno_id" in {
        row["name"] for row in query("PRAGMA table_info(documents)")
    }

    dossier = generate_contract_dossier({
        "contract": dict(query("SELECT * FROM contracts WHERE id=?", (contract_id,))[0]),
        "effective": {
            "original_start_date": "2026-08-01",
            "original_end_date": "2027-08-01",
            "current_start_date": "2027-08-01",
            "current_end_date": "2028-08-01",
            "current_value": 100000,
            "original_value": 100000,
            "current_instrument": "1º TERMO ADITIVO",
            "remaining_value": 100000,
            "lifecycle_status": "VIGENTE",
        },
        "amendments": [
            dict(query("SELECT * FROM amendments WHERE id=?", (amendment_id,))[0])
        ],
        "positions": [
            {
                **dict(query(
                    "SELECT * FROM contract_positions WHERE id=?", (position_id,)
                )[0]),
                "benefits": [],
            }
        ],
        "arts": [dict(query("SELECT * FROM arts WHERE id=?", (art_id,))[0])],
        "cnos": [dict(query(
            "SELECT * FROM contract_cnos WHERE id=?", (cno_id,)
        )[0])],
    })
    document = Document(BytesIO(dossier))
    text = "\n".join(
        paragraph.text for paragraph in document.paragraphs
    ) + "\n" + "\n".join(
        cell.text for table in document.tables for row in table.rows for cell in row.cells
    )
    assert "Engenheiro Civil" in text
    assert "90.000.00000/00" in text
    assert all(section.page_height > section.page_width for section in document.sections)
    assert query("PRAGMA integrity_check")[0][0] == "ok"
    print("Testes da versão 16 concluídos com sucesso.")


if __name__ == "__main__":
    run()
