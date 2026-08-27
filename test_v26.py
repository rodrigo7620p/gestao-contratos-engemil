import os
import tempfile
from datetime import date, timedelta
from io import BytesIO
from pathlib import Path

from docx import Document


def run():
    with tempfile.TemporaryDirectory() as directory:
        os.environ["GESTAO_DB_PATH"] = str(Path(directory) / "v26.db")
        os.environ["GESTAO_UPLOAD_DIR"] = str(Path(directory) / "uploads")

        from alerts import guarantee_expiry_notification_text
        from db import execute, init_db, query
        from guarantees import (
            calculate_required_amount,
            coverage_gap,
            guarantee_issues,
            operational_status,
        )
        from reports import generate_contract_dossier

        init_db()
        contract_id = execute(
            """INSERT INTO contracts(
            cost_center,client,contract_number,start_date,end_date,
            original_value,current_value,engineer_name,engineer_email)
            VALUES(?,?,?,?,?,?,?,?,?)""",
            (
                "01.02.00253",
                "AGÊNCIA BRASILEIRA DE PROMOÇÃO INTERNACIONAL DO TURISMO - EMBRATUR",
                "02/2024",
                "2026-01-01",
                "2027-12-31",
                1_000_000,
                1_000_000,
                "ENGENHEIRO TESTE",
                "engenheiro@engemil.com.br",
            ),
        )
        amendment_id = execute(
            """INSERT INTO amendments(
            contract_id,ordinal,kind,start_date,end_date,value)
            VALUES(?,'1º','TERMO ADITIVO','2026-07-01','2027-12-31',1100000)""",
            (contract_id,),
        )
        additional = calculate_required_amount(
            "GARANTIA_ADICIONAL",
            estimated_budget=1_000_000,
            proposal_value=780_000,
        )
        assert additional == 220_000
        percentage = calculate_required_amount(
            "PERCENTUAL_BASE", calculation_base=1_100_000, percentage=5
        )
        assert percentage == 55_000
        guarantee_id = execute(
            """INSERT INTO contract_guarantees(
            contract_id,amendment_id,guarantee_type,instrument_scope,modality,
            legal_basis,calculation_method,calculation_base,percentage,
            required_amount,guaranteed_amount,provider_name,policy_number,
            susep_registration,start_date,end_date,request_status,
            responsible_email,notification_enabled)
            VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,1)""",
            (
                contract_id,
                amendment_id,
                "GARANTIA CONTRATUAL",
                "ADITIVO / INSTRUMENTO",
                "SEGURO-GARANTIA",
                "Lei nº 14.133/2021, arts. 96 a 102",
                "PERCENTUAL_BASE",
                1_100_000,
                5,
                percentage,
                47_460.49,
                "JUNTO SEGUROS S.A.",
                "04-0775-0526688",
                "054362026000407750526688",
                date.today().isoformat(),
                (date.today() + timedelta(days=30)).isoformat(),
                "ACEITA",
                "engenheiro@engemil.com.br",
            ),
        )
        coverage_id = execute(
            """INSERT INTO guarantee_coverages(
            guarantee_id,coverage_name,insured_limit,start_date,end_date,deductible)
            VALUES(?,?,?,?,?,?)""",
            (
                guarantee_id,
                "Trabalhista e Previdenciária",
                47_460.49,
                date.today().isoformat(),
                (date.today() + timedelta(days=30)).isoformat(),
                "Sem franquia",
            ),
        )
        endorsement_id = execute(
            """INSERT INTO guarantee_endorsements(
            guarantee_id,endorsement_number,movement_type,new_end_date,new_amount,
            request_status) VALUES(?,?,?,?,?,'SOLICITADA')""",
            (
                guarantee_id,
                "END-001",
                "ENDOSSO",
                (date.today() + timedelta(days=365)).isoformat(),
                55_000,
            ),
        )
        assert coverage_id and endorsement_id
        assert coverage_gap(55_000, 47_460.49) == 7_539.51
        guarantee = dict(query(
            "SELECT * FROM contract_guarantees WHERE id=?", (guarantee_id,)
        )[0])
        assert operational_status(guarantee) == "A VENCER"
        issues = guarantee_issues(guarantee, "2027-12-31")
        assert "valor garantido inferior ao valor exigido" not in issues
        assert "vigência da garantia termina antes da vigência contratual" in issues

        document_columns = {
            row["name"] for row in query("PRAGMA table_info(documents)")
        }
        assert {"guarantee_id", "guarantee_endorsement_id"} <= document_columns
        assert query(
            "SELECT name FROM sqlite_master WHERE type='table' AND name='guarantee_coverages'"
        )

        alert_row = {
            **guarantee,
            "cost_center": "01.02.00253",
            "client": "AGÊNCIA BRASILEIRA DE PROMOÇÃO INTERNACIONAL DO TURISMO - EMBRATUR",
            "contract_number": "02/2024",
        }
        subject, body = guarantee_expiry_notification_text(alert_row, date.today())
        assert subject == "01.02.00253_EMBRATUR_[GARANTIA VENCE EM 30 DIA(S)]"
        assert "R$ 55.000,00" in body
        assert "Valor garantido" not in body
        assert "documentos comprobatórios" in body

        payload = {
            "contract": dict(query("SELECT * FROM contracts WHERE id=?", (contract_id,))[0]),
            "effective": {
                "current_instrument": "1º Termo Aditivo",
                "current_start_date": "2026-07-01",
                "current_end_date": "2027-12-31",
                "current_value": 1_100_000,
                "original_start_date": "2026-01-01",
                "original_end_date": "2027-12-31",
                "original_value": 1_000_000,
            },
            "guarantees": [{
                **guarantee,
                "display_type": "GARANTIA CONTRATUAL",
                "instrument_reference": "1º TERMO ADITIVO",
                "operational_status": operational_status(guarantee),
                "coverage_gap": coverage_gap(55_000, 47_460.49),
                "coverages": [dict(query(
                    "SELECT * FROM guarantee_coverages WHERE id=?", (coverage_id,)
                )[0])],
                "endorsements": [dict(query(
                    "SELECT * FROM guarantee_endorsements WHERE id=?", (endorsement_id,)
                )[0])],
            }],
        }
        docx = generate_contract_dossier(payload)
        document = Document(BytesIO(docx))
        text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        text += "\n" + "\n".join(
            cell.text for table in document.tables for row in table.rows for cell in row.cells
        )
        assert "Garantias contratuais e seguros" in text
        assert "Trabalhista e Previdenciária" in text
        assert "END-001" in text

        app_source = Path("app.py").read_text(encoding="utf-8")
        assert 'APP_VERSION = "27"' in app_source
        assert '"Garantias e seguros"' in app_source
        assert "Enviar alertas automáticos de vigência em 60, 30 e 15 dias" in app_source

    print("Testes da versão 26 concluídos com sucesso.")


if __name__ == "__main__":
    run()
