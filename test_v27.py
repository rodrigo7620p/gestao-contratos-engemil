import os
import tempfile
from io import BytesIO
from pathlib import Path

from docx import Document

from art_management import art_number_key, organize_art_rows, professional_profiles


def run():
    sample_arts = [
        {
            "id": 1,
            "professional_name": "MATHEUS ANTONIO MILITAO DE MENEZES",
            "professional_title": "Engenheiro Civil",
            "professional_registration": "13814/D-DF",
            "art_number": "0720240009575",
        },
        {
            "id": 2,
            "professional_name": "REGITON QUEIROZ DE MENEZES",
            "professional_title": "Engenheiro Eletricista",
            "professional_registration": "2454/D-DF",
            "art_number": "0720240011286",
        },
        {
            "id": 3,
            "professional_name": "IURE ARAUJO SANTIAGO",
            "professional_title": "Engenheiro Mecânico",
            "professional_registration": "4224/D-GO",
            "art_number": "0720240011312",
        },
        {
            "id": 4,
            "professional_name": "MATHEUS ANTONIO MILITÃO DE MENEZES",
            "professional_title": "Engenheiro Civil",
            "professional_registration": "13814/D-DF",
            "art_number": "0720240093957",
        },
    ]
    organized = organize_art_rows(sample_arts)
    assert [item["art_number"] for item in organized] == [
        "0720240009575",
        "0720240093957",
        "0720240011286",
        "0720240011312",
    ]
    profiles = professional_profiles(sample_arts)
    assert len(profiles) == 3
    assert profiles[0]["professional_name"] == "MATHEUS ANTONIO MILITÃO DE MENEZES"
    assert art_number_key("ART 07.2024-00093957") == "ART07202400093957"

    with tempfile.TemporaryDirectory() as directory:
        os.environ["GESTAO_DB_PATH"] = str(Path(directory) / "v27.db")
        os.environ["GESTAO_UPLOAD_DIR"] = str(Path(directory) / "uploads")

        from db import execute, init_db, query
        from reports import generate_contract_dossier

        init_db()
        assert query(
            """SELECT name FROM sqlite_master
            WHERE type='table' AND name='contract_budget_dates'"""
        )
        contract_id = execute(
            """INSERT INTO contracts(cost_center,client,contract_number,current_value)
            VALUES(?,?,?,?)""",
            ("01.02.00999", "ÓRGÃO DE TESTE - ODT", "01/2026", 100_000),
        )
        execute(
            """INSERT INTO contract_budget_dates(
            contract_id,reference_date,description,notes) VALUES(?,?,?,?)""",
            (
                contract_id,
                "2025-11-30",
                "Orçamento contratual inicial",
                "Referência da licitação",
            ),
        )
        budget_dates = [
            dict(row) for row in query(
                "SELECT * FROM contract_budget_dates WHERE contract_id=?",
                (contract_id,),
            )
        ]
        payload = {
            "contract": dict(query(
                "SELECT * FROM contracts WHERE id=?", (contract_id,)
            )[0]),
            "effective": {"current_value": 100_000},
            "budget_dates": budget_dates,
            "guarantees": [{
                "guarantee_type": "GARANTIA CONTRATUAL",
                "display_type": "GARANTIA CONTRATUAL",
                "instrument_reference": "Contrato inicial",
                "modality": "SEGURO-GARANTIA",
                "policy_number": "TESTE-001",
                "required_amount": 5_000,
                "premium_value": 250,
                "start_date": "2026-01-01",
                "end_date": "2027-01-01",
                "operational_status": "VIGENTE",
            }],
        }
        document = Document(BytesIO(generate_contract_dossier(payload)))
        text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        text += "\n" + "\n".join(
            cell.text for table in document.tables for row in table.rows for cell in row.cells
        )
        assert "Data(s) do orçamento" in text
        assert "30/11/2025" in text
        assert "Orçamento contratual inicial" in text
        assert "Garantido/LMG" not in text
        assert "Diferença" not in text

    app_source = Path("app.py").read_text(encoding="utf-8")
    report_source = Path("reports.py").read_text(encoding="utf-8")
    alert_source = Path("alerts.py").read_text(encoding="utf-8")
    assert 'APP_VERSION = "27"' in app_source
    for removed_label in (
        "Orçamento estimado",
        "Valor da proposta",
        "Valor efetivamente garantido / LMG",
        "Vencimento da garantia",
        "Índice de reajuste",
    ):
        assert removed_label not in app_source
    assert "Valor garantido:" not in alert_source
    assert "Garantido/LMG" not in report_source
    assert "organize_art_rows" in app_source
    assert "contract_budget_dates" in app_source

    print("Testes da versão 27 concluídos com sucesso.")


if __name__ == "__main__":
    run()
